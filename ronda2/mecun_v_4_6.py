# -*- coding: utf-8 -*-
# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║           PIPELINE MECUN v4.6                                               ║
# ║                                                                              ║
# ║  MEJORAS SOBRE v4.5:                                                        ║
# ║  • FIX-24: Clasificación por categoria_id numérico                          ║
# ║  • FIX-25: Descubrimiento de categorías más estricto                        ║
# ║  • FIX-26: Dedupe semántico simple de categorías nuevas                     ║
# ║  • FIX-27: Límite máximo de categorías dinámicas                            ║
# ║  • FIX-28: Núcleo estricto y cross-eje con tolerancia controlada            ║
# ║                                                                              ║
# ║  INVARIANTE SAGRADA: ninguna propuesta es modificada en ningún punto        ║
# ╚══════════════════════════════════════════════════════════════════════════════╝


# ─────────────────────────────────────────────────────────────────────────────
# 0.  INSTALACIÓN DE DEPENDENCIAS
# ─────────────────────────────────────────────────────────────────────────────
import sys as _sys


def _en_notebook() -> bool:
    try:
        shell = get_ipython().__class__.__name__   # type: ignore[name-defined]
        return shell in ('ZMQInteractiveShell', 'Shell')
    except NameError:
        return False


if _en_notebook():
    import subprocess
    subprocess.run(
        [_sys.executable, '-m', 'pip', 'install', '-q',
         'python-docx', 'pandas', 'google-genai', 'tqdm'],
        check=False,
    )


# ─────────────────────────────────────────────────────────────────────────────
# 1.  IMPORTS
# ─────────────────────────────────────────────────────────────────────────────
import re
import time
import json
import pickle
import unicodedata
from pathlib import Path
from datetime import datetime
from difflib import SequenceMatcher
from typing import Any, Callable

import pandas as pd
from google import genai
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tqdm import tqdm


# ─────────────────────────────────────────────────────────────────────────────
# 2.  CONFIGURACIÓN  ✏️  ← EDITAR ESTAS VARIABLES ANTES DE EJECUTAR
# ─────────────────────────────────────────────────────────────────────────────

# 🔑  Clave de API de Google Gemini
API_KEY = ''

# 📂  Rutas de Google Drive
DATA_SOURCE_PATH  = ''
REPORTES_BASE_DIR = Path('/content/drive/MyDrive/Reportes_Generados')
CHECKPOINT_DIR    = Path('/content/drive/MyDrive/MECUN_Checkpoints')

# 🤖  Modelos Gemini
MODEL_FLASH = 'gemini-2.5-flash'
MODEL_PRO   = 'gemini-2.5-pro'

# ⚙️  Parámetros del pipeline
MARGEN_ERROR_SLOVIN  = 0.10
UMBRAL_CONFIANZA     = 70
TAMANO_LOTE_CLASIF   = 10
TAMANO_LOTE_AGRUPAR  = 50
PAUSA_ENTRE_LLAMADAS = 2
UMBRAL_PROPAGACION   = 60
CONFIANZA_CROSS_EJE  = 90

# 🛡️  Calidad del pipeline
MODO_ESTRICTO_NUCLEO      = True
MODO_ESTRICTO_CROSS_EJE   = False
TASA_MAX_FALLOS_CROSS_EJE = 0.15
DEBUG_JSON                = False

# 🏷️  Estrictez en categorías dinámicas
USAR_DESCUBRIMIENTO_DINAMICO = True
MAX_CATEGORIAS_DINAMICAS     = 12
SIMILITUD_MAX_CATEGORIAS     = 0.88
LONGITUD_MIN_CATEGORIA       = 12

# 💾  Checkpoints
FORZAR_RECALCULO_PASS1 = False


# ─────────────────────────────────────────────────────────────────────────────
# Validación y setup de rutas
# ─────────────────────────────────────────────────────────────────────────────
if not API_KEY or API_KEY == 'PEGA_AQUI_TU_GEMINI_API_KEY':
    raise ValueError(
        'Debes pegar tu GEMINI API KEY en la variable API_KEY antes de ejecutar.'
    )

REPORTES_BASE_DIR.mkdir(parents=True, exist_ok=True)
CHECKPOINT_DIR.mkdir(parents=True, exist_ok=True)

PASS1_CHECKPOINT = CHECKPOINT_DIR / 'pass1_cross_eje.pkl'
client           = genai.Client(api_key=API_KEY)


# ─────────────────────────────────────────────────────────────────────────────
# Ejes temáticos y diccionario de categorías
# ─────────────────────────────────────────────────────────────────────────────
EJE_ACTUAL = 'Arquitectura del (co)gobierno universitario.'

EJES_VALIDOS = {
    'EJ1': 'Arquitectura del (co)gobierno universitario.',
    'EJ2': 'Elección y designación de autoridades académicas.',
    'EJ3': 'Formas y mecanismos de participación democrática.',
    'EJ4': 'Fortalecimiento de la cultura política democrática de la comunidad universitaria.',
    'EJ5': 'Formas de organización y tejido social universitario.',
}
NOMBRES_EJES = list(EJES_VALIDOS.values())

DICCIONARIO_CATEGORIAS = {
    'Formas y mecanismos de participación democrática.': [
        'Composición de Cuerpos Colegiados y Representación Estamental',
        'Elección, Designación y Revocatoria de Autoridades y Representantes',
        'Transparencia, Rendición de Cuentas y Control Social',
        'Participación y Articulación del Estamento de Egresados',
        'Espacios de Participación de Base, Organización Autónoma y Deliberación',
        'Marco Normativo de la Participación y Reforma Estatutaria',
        'Mecanismos Directos y Vinculantes de Decisión Universitaria',
        'Participación Presupuestal, Logística y Apoyo a la Participación',
        'Comunicación, Formación y Cultura para la Participación Democrática',
        'Casos Aislados / Otros',
    ],
    'Arquitectura del (co)gobierno universitario.': [
        'Diseño, Composición y Funciones de los Cuerpos Colegiados',
        'Mecanismos Electorales, Designación y Revocatoria de Autoridades',
        'Descentralización, Autonomía y Representación de Sedes y Unidades',
        'Participación Amplia, Inclusiva y Deliberativa de la Comunidad Universitaria',
        'Transparencia, Veeduría y Rendición de Cuentas en el Cogobierno',
        'Criterios de Idoneidad y Perfiles para Cargos Directivos',
        'Marco Normativo del Cogobierno y Procesos de Reforma Institucional',
        'Casos Aislados / Otros',
    ],
    'Fortalecimiento de la cultura política democrática de la comunidad universitaria.': [
        'Gobernanza Democrática, Participación y Control Institucional',
        'Formación Política, Pensamiento Crítico y Cultura Democrática',
        'Inclusión, Diversidad, Equidad y Enfoques Diferenciales',
        'Bienestar Integral, Salud Mental y Políticas del Cuidado',
        'Transformación Curricular y Vinculación Social de la Universidad',
        'Seguridad, Convivencia y Derechos Humanos en la Vida Universitaria',
        'Gestión del Campus, Espacios Universitarios y Sostenibilidad',
        'Apoyo, Reconocimiento y Fortalecimiento de Organizaciones Universitarias',
        'Articulación Integral de Políticas y Estrategias Institucionales',
        'Casos Aislados / Otros',
    ],
    'Elección y designación de autoridades académicas.': [
        'Mecanismos de Elección y Designación de Autoridades Académicas',
        'Ponderación del Voto y Representación Estamental en las Elecciones',
        'Reglas del Proceso Electoral, Escrutinio y Voto en Blanco',
        'Diseño y Composición de Órganos Colegiados y Electorales',
        'Transparencia, Control, Rendición de Cuentas y Revocatoria de Mandato',
        'Criterios de Idoneidad, Requisitos y Habilitación de Candidaturas',
        'Descentralización, Autonomía y Representación Territorial de Sedes',
        'Participación, Comunicación y Cultura Democrática en los Procesos Electorales',
        'Participación y Descentralización Presupuestal en la Gobernanza',
        'Casos Aislados / Otros',
    ],
    'Formas de organización y tejido social universitario.': [
        'Gestión Ambiental, Biodiversidad y Sostenibilidad del Campus',
        'Fortalecimiento, Recursos y Espacios para Organizaciones Universitarias',
        'Reestructuración Académica, Epistemológica e Interdisciplinaria',
        'Condiciones Laborales, Formalización y Empleabilidad de la Comunidad Universitaria',
        'Gestión Operativa, Infraestructura y Servicios Universitarios',
        'Propuestas Indefinidas, Remisivas o Sin Contenido Explícito',
        'Casos Aislados / Otros',
    ],
}
CATEGORIAS_BASE = DICCIONARIO_CATEGORIAS.get(EJE_ACTUAL, [])

# 🎨  Colores del documento Word
COLOR_AZUL_OSCURO = '1F3864'
COLOR_AZUL_MEDIO  = '2E5FA3'
COLOR_GRIS_CLARO  = 'F2F2F2'
COLOR_VERDE       = '1E7B34'
COLOR_NARANJA     = 'B8720D'
COLOR_ROJO        = 'C00000'
COLOR_LILA        = '5B2C8D'


# ─────────────────────────────────────────────────────────────────────────────
# 3.  FUNCIONES DE SOPORTE GENERAL
# ─────────────────────────────────────────────────────────────────────────────

def crear_carpeta_salida(eje: str) -> Path:
    slug = re.sub(r'[^a-z0-9]', '_', eje.lower())[:30].strip('_')
    carpeta = REPORTES_BASE_DIR / slug
    carpeta.mkdir(parents=True, exist_ok=True)
    return carpeta


# Timestamps y carpeta de salida
TS              = datetime.now().strftime('%Y%m%d_%H%M%S')
FECHA_EJECUCION = datetime.now().strftime('%d/%m/%Y %H:%M')
CARPETA_SALIDA  = crear_carpeta_salida(EJE_ACTUAL)


def _safe_str(val: Any, default: str = '') -> str:
    """Devuelve str(val) o default si val es None / NaN / vacío."""
    if val is None:
        return default
    try:
        if pd.isna(val):
            return default
    except (TypeError, ValueError):
        pass
    s = str(val).strip()
    return s if s not in ('nan', 'NaN', 'None', 'null') else default


def leer_csv_robusto(ruta: str) -> pd.DataFrame:
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
    for enc in encodings:
        try:
            df = pd.read_csv(ruta, encoding=enc, on_bad_lines='skip')
            print(f'  CSV cargado con encoding: {enc} ({len(df)} filas)')
            return df
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f'  ⚠ Error con {enc}: {e}')
            continue
    raise ValueError(
        f'No se pudo leer "{ruta}" con ninguno de los encodings: {encodings}'
    )


def calcular_muestra_estadistica(total: int, error: float) -> int:
    if total <= 0:
        return 0
    n = int(total / (1 + total * (error ** 2)))
    return max(min(n, 200), 5)


def llamar_api(prompt: str, model: str = MODEL_FLASH,
               max_reintentos: int = 3) -> str | None:
    for intento in range(1, max_reintentos + 1):
        try:
            respuesta = client.models.generate_content(
                model=model,
                contents=prompt,
            )
            return respuesta.text.strip() if respuesta and getattr(respuesta, 'text', None) else None
        except Exception as e:
            msg = str(e).lower()
            espera = 20 * intento if ('429' in msg or 'quota' in msg) else 3 * intento
            print(f'  [⏳] Intento {intento}/{max_reintentos}. Pausa {espera}s. Error: {e}')
            time.sleep(espera)
    return None


def _debug_respuesta_cruda(contexto: str, respuesta: str | None) -> None:
    if not DEBUG_JSON:
        return
    if not respuesta:
        print(f'  [DEBUG {contexto}] Respuesta vacía/None')
        return
    preview = respuesta[:800].replace('\n', '\\n')
    print(f'  [DEBUG {contexto}] Preview respuesta: {preview}')


def _validar_items_clasificacion_por_id(items: Any) -> bool:
    if not isinstance(items, list):
        return False
    for x in items:
        if not isinstance(x, dict):
            return False
        if 'id' not in x or 'categoria_id' not in x or 'confianza' not in x:
            return False
    return True


def _validar_items_cross_eje(items: Any) -> bool:
    if not isinstance(items, list):
        return False
    for x in items:
        if not isinstance(x, dict):
            return False
        if 'eje' not in x or 'razon' not in x or 'confianza' not in x:
            return False
    return True


def _validar_items_agrupacion(items: Any) -> bool:
    if not isinstance(items, list):
        return False
    for x in items:
        if not isinstance(x, dict):
            return False
        if 'sintesis' not in x or 'ids' not in x:
            return False
        if not isinstance(x['ids'], list):
            return False
    return True


def _parsear_json_respuesta(
    respuesta: str | None,
    lote_ids: list,
    contexto: str = '',
    validador: Callable[[Any], bool] | None = None
) -> tuple[list[Any], bool]:
    if not respuesta:
        print(f'  [⚠ JSON] Respuesta vacía/None. {contexto}')
        return [], False

    limpia = re.sub(r'```(?:json)?', '', respuesta).strip()

    def _es_valido(items: Any) -> bool:
        if validador is None:
            return isinstance(items, list)
        return validador(items)

    try:
        items = json.loads(limpia)
        if _es_valido(items):
            return items, True
    except json.JSONDecodeError:
        pass

    match = re.search(r'\[.*\]', limpia, re.DOTALL)
    if match:
        try:
            items = json.loads(match.group())
            if _es_valido(items):
                print(f'  [ℹ JSON] Extracción regex exitosa. {contexto}')
                return items, True
        except json.JSONDecodeError:
            pass

    n_ids = len(lote_ids) if lote_ids else '?'
    print(
        f'  [⚠ JSON] Parse fallido ({contexto}, {n_ids} propuestas) '
        '→ estructura inválida.'
    )
    _debug_respuesta_cruda(contexto, respuesta)
    return [], False


def _limpiar_parrafo_compat(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def _deduplicar_preservando_orden(seq: list[str]) -> list[str]:
    vistos = set()
    out = []
    for x in seq:
        if x not in vistos:
            vistos.add(x)
            out.append(x)
    return out


def _normalizar_ids_validos(ids_raw: Any, mapeo: dict[str, Any]) -> list[str]:
    if not isinstance(ids_raw, list):
        return []
    normalizados = []
    for x in ids_raw:
        sx = str(x).strip()
        if sx in mapeo:
            normalizados.append(sx)
    return _deduplicar_preservando_orden(normalizados)


def _normalizar_etiqueta_categoria(txt: str) -> str:
    txt = re.sub(r'\s+', ' ', txt.strip())
    txt = txt.strip(' .;:-')
    return txt


def _parecido(a: str, b: str) -> float:
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def _filtrar_categorias_dinamicas(categorias_raw: list[str],
                                   categorias_base: list[str]) -> list[str]:
    resultado: list[str] = []
    universo_existente = [c for c in categorias_base if 'Casos Aislados' not in c]

    for cat in categorias_raw:
        c = _normalizar_etiqueta_categoria(cat)

        if len(c) < LONGITUD_MIN_CATEGORIA:
            continue

        c_low = c.lower()
        if c_low in {'ninguna', 'otros', 'otras', 'varios', 'miscelánea', 'miscelanea'}:
            continue

        if any(_parecido(c, base) >= SIMILITUD_MAX_CATEGORIAS for base in universo_existente):
            continue

        if any(_parecido(c, ya) >= SIMILITUD_MAX_CATEGORIAS for ya in resultado):
            continue

        resultado.append(c)

        if len(resultado) >= MAX_CATEGORIAS_DINAMICAS:
            break

    return resultado


def _validar_cobertura_categoria(categoria: str,
                                  mapeo: dict[str, Any],
                                  grupos_finales: list[dict]) -> None:
    esperados = set(mapeo.keys())
    conteos: dict[str, int] = {k: 0 for k in esperados}

    for grupo in grupos_finales:
        ids_g = grupo.get('ids', [])
        if not isinstance(ids_g, list):
            raise RuntimeError(f'Categoría "{categoria}": grupo con ids inválidos.')
        for x in ids_g:
            if x not in conteos:
                raise RuntimeError(f'Categoría "{categoria}": id fuera de mapeo: {x}')
            conteos[x] += 1

    faltantes  = [k for k, v in conteos.items() if v == 0]
    duplicados = [k for k, v in conteos.items() if v > 1]

    if faltantes or duplicados:
        raise RuntimeError(
            f'Cobertura inválida en categoría "{categoria}". '
            f'Faltantes={faltantes[:10]} Duplicados={duplicados[:10]}'
        )


# ═════════════════════════════════════════════════════════════════════════════
# MÓDULO A: PROPAGACIÓN DE LINKS DE RELATORÍA
# ═════════════════════════════════════════════════════════════════════════════

def _normalizar_claustro(nombre: str) -> str:
    if pd.isna(nombre):
        return ''
    return re.sub(r'\s+', ' ', str(nombre).lower().strip())


def _score_doc_match(r_sin: pd.Series, r_con: pd.Series) -> tuple[int, list[str]]:
    score = 0
    reasons: list[str] = []

    c1 = _normalizar_claustro(r_sin.get('claustro_norm', ''))
    c2 = _normalizar_claustro(r_con.get('claustro_norm', ''))
    if c1 and c2:
        if c1 == c2:
            score += 50
            reasons.append('exact_claustro(+50)')
        else:
            ratio = SequenceMatcher(None, c1, c2).ratio()
            if ratio >= 0.85:
                pts = int(ratio * 40)
                score += pts
                reasons.append(f'fuzzy_{ratio:.2f}(+{pts})')

    e1 = str(r_sin.get('email_norm', '')).lower().strip()
    e2 = str(r_con.get('email_norm', '')).lower().strip()
    if e1 and e2 and e1 == e2:
        score += 35
        reasons.append('exact_email(+35)')

    t1 = r_sin.get('ts_parsed')
    t2 = r_con.get('ts_parsed')
    if pd.notna(t1) and pd.notna(t2):
        diff_h = abs((t1 - t2).total_seconds()) / 3600
        if diff_h <= 2:
            score += 10
            reasons.append('<2h(+10)')
        elif diff_h <= 24:
            score += 5
            reasons.append('<24h(+5)')

    return score, reasons


def propagar_links_relatorias(df_raw: pd.DataFrame,
                               col_doc: str | None,
                               col_claustro: str | None,
                               col_email: str | None,
                               col_ts: str | None) -> pd.DataFrame:
    df = df_raw.copy()

    if col_doc is None or col_claustro is None:
        print('  ⚠ col_doc o col_claustro no encontradas — propagación omitida.')
        df['claustro_norm']       = ''
        df['email_norm']          = ''
        df['ts_parsed']           = pd.NaT
        df['Documento_Propagado'] = None
        df['Doc_Score']           = 0
        df['Doc_Fuente']          = 'sin_doc'
        df['Doc_Razones']         = ''
        return df

    df['claustro_norm'] = df[col_claustro].apply(_normalizar_claustro)
    df['email_norm']    = df[col_email].str.lower().str.strip() if col_email else ''
    df['ts_parsed']     = (
        pd.to_datetime(df[col_ts], dayfirst=True, errors='coerce')
        if col_ts else pd.NaT
    )
    df['Documento_Propagado'] = df[col_doc]
    df['Doc_Score']   = df[col_doc].apply(
        lambda x: 95 if pd.notna(x) and str(x).strip() else 0
    )
    df['Doc_Fuente']  = df[col_doc].apply(
        lambda x: 'original' if pd.notna(x) and str(x).strip() else 'sin_doc'
    )
    df['Doc_Razones'] = df[col_doc].apply(
        lambda x: 'link subido directamente' if pd.notna(x) and str(x).strip() else ''
    )

    filas_con_doc = df[df['Doc_Fuente'] == 'original']
    for idx in df[df['Doc_Fuente'] == 'sin_doc'].index:
        row_sin = df.loc[idx]
        mejor_score, mejor_doc, mejor_razones = 0, None, []
        for _, row_con in filas_con_doc.iterrows():
            s, r = _score_doc_match(row_sin, row_con)
            if s > mejor_score:
                mejor_score, mejor_doc, mejor_razones = s, row_con[col_doc], r
        if mejor_score >= UMBRAL_PROPAGACION and mejor_doc:
            df.at[idx, 'Documento_Propagado'] = mejor_doc
            df.at[idx, 'Doc_Score']           = mejor_score
            df.at[idx, 'Doc_Fuente']          = 'propagado'
            df.at[idx, 'Doc_Razones']         = ' | '.join(mejor_razones)

    total_orig = (df['Doc_Fuente'] == 'original').sum()
    total_prop = (df['Doc_Fuente'] == 'propagado').sum()
    total_sin  = (df['Doc_Fuente'] == 'sin_doc').sum()
    print(f'  → Links: {total_orig} originales | {total_prop} propagados | {total_sin} sin doc')
    return df


# ═════════════════════════════════════════════════════════════════════════════
# MÓDULO B: DETECCIÓN Y COPIA CROSS-EJE
# ═════════════════════════════════════════════════════════════════════════════

DESCRIPCION_EJES = {
    'Arquitectura del (co)gobierno universitario.': (
        'Cuerpos colegiados, consejos directivos, composición y funciones de instancias '
        'de decisión estatutaria (CSU, CA, Consejos de Sede/Facultad).'
    ),
    'Elección y designación de autoridades académicas.': (
        'Procedimientos para elegir o designar rector, decanos, directores de sede; '
        'mecanismos electorales, requisitos, consultas, voto universal.'
    ),
    'Formas y mecanismos de participación democrática.': (
        'Instrumentos democráticos: plebiscitos, referendos, cabildos, asambleas, '
        'veedurías, plataformas digitales de voto, presupuestos participativos.'
    ),
    'Fortalecimiento de la cultura política democrática de la comunidad universitaria.': (
        'Formación ciudadana, cátedras, deliberación pública, ética, transparencia, '
        'garantías para movilización y ejercicio político estudiantil/docente.'
    ),
    'Formas de organización y tejido social universitario.': (
        'Colectivos estudiantiles, sindicatos, asociaciones, grupos culturales/artísticos, '
        'diversidades, economía solidaria, cohesión social orgánica en los campus.'
    ),
}

PROMPT_CROSS_EJE = """
Eres un analista normativo experto en la reforma universitaria de la UNAL (MECUN).
Determina si esta propuesta, presentada bajo un eje, también contiene ideas que
deben ser revisadas por las comisiones de OTROS ejes.

EJES TEMÁTICOS Y SU ALCANCE:
{ejes_desc}

PROPUESTA:
Eje original: {eje_original}
Texto: {texto}

INSTRUCCIONES:
- Solo copia a otro eje si hay contenido concreto y accionable para esa comisión.
- Si es 100% del eje original responde exactamente: SOLO_EJE_ORIGINAL
- Si toca otros ejes, responde en JSON válido, sin backticks, sin comentarios:
  [{{"eje": "Nombre exacto", "razon": "Qué parte es relevante", "confianza": 80}}]
- No agregues texto antes ni después del JSON.
"""


def _cfg_cross_eje() -> dict[str, Any]:
    return {
        'confianza_estricta': max(CONFIANZA_CROSS_EJE, 90),
        'max_candidatos': 2,
        'max_chars_texto': 1600,
        'usar_fallback_legado_solo_error_tecnico': True,
        'min_score_candidato': 3,
        'min_score_evidencia': 3,
    }


def _get_descripcion_ejes_estricta() -> dict[str, str]:
    return {
        'Arquitectura del (co)gobierno universitario.': (
            'Órganos, cuerpos colegiados, composición, competencias, funciones, '
            'relaciones entre niveles de gobierno, representación estamental y '
            'estructura de cogobierno. Incluye CSU, Consejo Académico, Consejos de '
            'Sede y Consejos de Facultad. No incluye por sí solo reglas para elegir '
            'autoridades específicas.'
        ),
        'Elección y designación de autoridades académicas.': (
            'Reglas, procedimientos y requisitos para elegir o designar rectoría, '
            'decanaturas, direcciones de sede y otros cargos de autoridad académica. '
            'Incluye escrutinio, consulta, elegibilidad, candidaturas, revocatoria '
            'y mecanismos electorales ligados a autoridades. No cubre por sí solo '
            'la estructura general de órganos colegiados.'
        ),
        'Formas y mecanismos de participación democrática.': (
            'Mecanismos concretos, procedimentales o vinculantes de participación y '
            'decisión colectiva: asambleas, cabildos, consultas, referendos, '
            'plebiscitos, veedurías, presupuestos participativos, votación directa '
            'o electrónica e iniciativas normativas. No equivale a cualquier mención '
            'genérica a participación.'
        ),
        'Fortalecimiento de la cultura política democrática de la comunidad universitaria.': (
            'Formación política, pedagogía democrática, deliberación, pensamiento '
            'crítico, convivencia, ética pública, garantías para el ejercicio político, '
            'derechos humanos, cuidado y condiciones culturales para una vida '
            'democrática universitaria. No equivale a cualquier mención genérica a democracia.'
        ),
        'Formas de organización y tejido social universitario.': (
            'Colectivos, asociaciones, sindicatos, redes, organizaciones estudiantiles, '
            'docentes, de egresados o administrativas, y su fortalecimiento material, '
            'reconocimiento, articulación, autonomía y espacios. No equivale a '
            'cualquier mención genérica a comunidad universitaria.'
        ),
    }


def _get_lexicon_ejes() -> dict[str, dict[str, list[str]]]:
    return {
        'Arquitectura del (co)gobierno universitario.': {
            'nuclear': [
                'cuerpo colegiado', 'cuerpos colegiados', 'consejo superior',
                'consejo academico', 'consejo de sede', 'consejo de facultad',
                'composicion', 'competencias', 'funciones', 'estructura de gobierno',
                'cogobierno', 'representacion estamental', 'instancia de decision'
            ],
            'support': [
                'estatuto general', 'gobernanza', 'organo colegiado',
                'arquitectura institucional'
            ],
        },
        'Elección y designación de autoridades académicas.': {
            'nuclear': [
                'rector', 'rectoria', 'decano', 'decanatura', 'vicerrector',
                'vicerrectoria', 'director de sede', 'direccion academica',
                'direccion academico administrativa', 'eleccion', 'designacion',
                'consulta', 'escrutinio', 'candidatura', 'revocatoria',
                'elegibilidad', 'requisitos para el cargo'
            ],
            'support': [
                'voto', 'terna', 'habilitacion', 'proceso electoral',
                'autoridad academica'
            ],
        },
        'Formas y mecanismos de participación democrática.': {
            'nuclear': [
                'cabildo', 'asamblea', 'consulta', 'referendo', 'plebiscito',
                'veeduria', 'votacion directa', 'votacion electronica',
                'mecanismo vinculante', 'presupuesto participativo',
                'iniciativa normativa', 'control social'
            ],
            'support': [
                'participacion vinculante', 'deliberacion', 'mecanismo de participacion',
                'decisiones colectivas'
            ],
        },
        'Fortalecimiento de la cultura política democrática de la comunidad universitaria.': {
            'nuclear': [
                'formacion politica', 'cultura democratica', 'pensamiento critico',
                'deliberacion publica', 'derechos humanos', 'convivencia',
                'salud mental', 'cuidado', 'etica publica', 'no estigmatizacion',
                'pedagogia democratica', 'formacion ciudadana'
            ],
            'support': [
                'bienestar', 'garantias para la participacion', 'clima democratico'
            ],
        },
        'Formas de organización y tejido social universitario.': {
            'nuclear': [
                'colectivo', 'colectivos', 'sindicato', 'sindicatos', 'asociacion',
                'asociaciones', 'organizacion estudiantil', 'organizaciones universitarias',
                'tejido social', 'redes estudiantiles', 'red organizativa',
                'egresados organizados'
            ],
            'support': [
                'espacios para organizaciones', 'financiacion organizativa',
                'autonomia organizativa', 'fortalecimiento organizativo'
            ],
        },
    }


def _norm_text(s: str) -> str:
    s = str(s or '').lower().strip()
    s = ''.join(
        c for c in unicodedata.normalize('NFD', s)
        if unicodedata.category(c) != 'Mn'
    )
    s = re.sub(r'[^a-z0-9\s]', ' ', s)
    return re.sub(r'\s+', ' ', s).strip()


def _score_eje_destino(texto: str, eje_dest: str) -> tuple[int, list[str], int]:
    lex = _get_lexicon_ejes()
    txt = _norm_text(texto)

    hits: list[str] = []
    score = 0
    nuclear_hits = 0

    for t in lex[eje_dest]['nuclear']:
        if t in txt:
            hits.append(t)
            score += 3
            nuclear_hits += 1

    for t in lex[eje_dest]['support']:
        if t in txt:
            hits.append(t)
            score += 1

    return score, hits, nuclear_hits


def _ejes_candidatos_cross(texto: str, eje_actual: str) -> list[str]:
    cfg = _cfg_cross_eje()
    candidatos: list[tuple[str, int, int]] = []

    for eje in NOMBRES_EJES:
        if eje == eje_actual:
            continue

        score, _, nuclear_hits = _score_eje_destino(texto, eje)
        if score >= cfg['min_score_candidato'] and nuclear_hits >= 1:
            candidatos.append((eje, score, nuclear_hits))

    candidatos.sort(key=lambda x: (x[1], x[2]), reverse=True)
    return [eje for eje, _, _ in candidatos[:cfg['max_candidatos']]]


def _evidencia_literal_valida(evidencia: str, texto_fuente: str, eje_dest: str) -> bool:
    cfg = _cfg_cross_eje()

    ev = _norm_text(evidencia)
    tx = _norm_text(texto_fuente)

    if len(ev) < 12:
        return False

    if ev not in tx:
        return False

    score_ev, _, nuclear_hits_ev = _score_eje_destino(ev, eje_dest)
    if score_ev < cfg['min_score_evidencia']:
        return False
    if nuclear_hits_ev < 1:
        return False

    return True


def _validar_items_cross_eje_estricto(items: Any) -> bool:
    if not isinstance(items, list):
        return False

    for x in items:
        if not isinstance(x, dict):
            return False
        if 'eje' not in x or 'razon' not in x or 'confianza' not in x or 'evidencia_textual' not in x:
            return False

    return True


def _construir_prompt_cross_eje_estricto(
    eje_actual: str,
    candidatos: list[str],
    texto_fuente: str
) -> str:
    descripciones = _get_descripcion_ejes_estricta()
    ejes_desc = '\n'.join(f'• {eje}: {descripciones[eje]}' for eje in candidatos)

    return f"""
Eres un filtro de ALTA PRECISIÓN para propuestas MECUN-UNAL.

Tu objetivo NO es maximizar copias entre ejes.
Tu objetivo es EVITAR FALSOS POSITIVOS.

EJE ORIGINAL:
{eje_actual}

SOLO puedes evaluar estos posibles ejes destino:
{ejes_desc}

REGLAS DURAS:
- No copies por lenguaje genérico como "democracia", "participación", "comunidad", "inclusión" o "representación" si no hay contenido nuclear del eje destino.
- Solo copia si hay un componente concreto, accionable y claramente competencia del eje destino.
- Debes incluir "evidencia_textual" como fragmento LITERAL tomado del texto.
- Si la propuesta pertenece solo al eje original, responde exactamente: SOLO_EJE_ORIGINAL
- Si dudas, responde: SOLO_EJE_ORIGINAL

TEXTO:
{texto_fuente}

Si aplica cross-eje, responde SOLO JSON válido:
[{{"eje":"Nombre exacto","evidencia_textual":"fragmento literal","razon":"breve","confianza":92}}]
""".strip()


def _resolver_cross_eje_estricto_para_texto(
    eje_actual: str,
    texto_fuente: str
) -> list[dict]:
    cfg = _cfg_cross_eje()

    candidatos = _ejes_candidatos_cross(texto_fuente, eje_actual)
    if not candidatos:
        return []

    prompt = _construir_prompt_cross_eje_estricto(
        eje_actual=eje_actual,
        candidatos=candidatos,
        texto_fuente=texto_fuente[:cfg['max_chars_texto']]
    )

    respuesta = llamar_api(prompt, model=MODEL_FLASH)
    if not respuesta:
        raise RuntimeError('cross_estricto_respuesta_vacia')

    if 'SOLO_EJE_ORIGINAL' in respuesta.upper():
        return []

    matches, ok = _parsear_json_respuesta(
        respuesta,
        [],
        'cross-eje-estricto',
        validador=_validar_items_cross_eje_estricto
    )
    if not ok:
        raise RuntimeError('cross_estricto_json_invalido')

    salidas: list[dict] = []

    for m in matches:
        eje_dest  = str(m.get('eje', '')).strip()
        razon     = str(m.get('razon', '')).strip()
        evidencia = str(m.get('evidencia_textual', '')).strip()

        try:
            confianza = int(m.get('confianza', 0))
        except (TypeError, ValueError):
            confianza = 0

        if eje_dest not in candidatos:
            continue
        if eje_dest == eje_actual:
            continue
        if confianza < cfg['confianza_estricta']:
            continue
        if not _evidencia_literal_valida(evidencia, texto_fuente, eje_dest):
            continue

        score_dest, _, nuclear_hits = _score_eje_destino(texto_fuente, eje_dest)
        if score_dest < cfg['min_score_candidato'] or nuclear_hits < 1:
            continue

        salidas.append({
            'eje': eje_dest,
            'razon': f'{razon} | evidencia: "{evidencia}"',
            'confianza': confianza,
        })

    salidas = [
        x for i, x in enumerate(salidas)
        if x['eje'] not in {y['eje'] for y in salidas[:i]}
    ]

    return salidas


def _resolver_cross_eje_legado_para_texto(
    eje_actual: str,
    texto: str
) -> list[dict]:
    ejes_desc_txt = '\n'.join(
        f'• {eje}: {desc}' for eje, desc in DESCRIPCION_EJES.items()
    )

    prompt = PROMPT_CROSS_EJE.format(
        ejes_desc=ejes_desc_txt,
        eje_original=eje_actual,
        texto=str(texto)[:800],
    )

    respuesta = llamar_api(prompt, model=MODEL_FLASH)
    if not respuesta:
        raise RuntimeError('cross_legado_respuesta_vacia')

    if 'SOLO_EJE_ORIGINAL' in respuesta.upper():
        return []

    matches, ok = _parsear_json_respuesta(
        respuesta,
        [],
        'cross-eje-legado',
        validador=_validar_items_cross_eje
    )
    if not ok:
        raise RuntimeError('cross_legado_json_invalido')

    salidas: list[dict] = []

    for m in matches:
        eje_dest = str(m.get('eje', '')).strip()
        razon    = str(m.get('razon', '')).strip()

        try:
            confianza = int(m.get('confianza', 0))
        except (TypeError, ValueError):
            confianza = 0

        if eje_dest not in NOMBRES_EJES or eje_dest == eje_actual:
            continue
        if confianza < CONFIANZA_CROSS_EJE:
            continue

        salidas.append({'eje': eje_dest, 'razon': razon, 'confianza': confianza})

    salidas = [
        x for i, x in enumerate(salidas)
        if x['eje'] not in {y['eje'] for y in salidas[:i]}
    ]

    return salidas


def detectar_cross_eje_lote(df_eje: pd.DataFrame,
                              eje_actual: str) -> tuple[pd.DataFrame, dict[str, pd.DataFrame]]:
    cfg = _cfg_cross_eje()

    df_eje = df_eje.copy()
    df_eje['Ejes_Adicionales'] = [[] for _ in range(len(df_eje))]
    df_eje['Es_Copia']         = False
    df_eje['Eje_Origen_Copia'] = ''

    copias_por_eje: dict[str, list[dict]] = {
        e: [] for e in NOMBRES_EJES if e != eje_actual
    }

    print(f'  Detectando cross-eje para {len(df_eje)} propuestas…')

    evaluadas_cross       = 0
    fallos_estrictos      = 0
    fallos_no_recuperados = 0
    activaciones_fallback = 0

    for idx, row in tqdm(df_eje.iterrows(), total=len(df_eje),
                          desc='  Cross-eje', leave=False):
        texto_contextual = str(
            row.get('Texto_Contextualizado') or row.get('Propuesta') or ''
        ).strip()

        if not texto_contextual:
            continue

        evaluadas_cross += 1
        matches: list[dict] = []

        try:
            matches = _resolver_cross_eje_estricto_para_texto(
                eje_actual=eje_actual,
                texto_fuente=texto_contextual
            )
        except Exception as e:
            fallos_estrictos += 1

            if cfg['usar_fallback_legado_solo_error_tecnico']:
                activaciones_fallback += 1
                try:
                    matches = _resolver_cross_eje_legado_para_texto(
                        eje_actual=eje_actual,
                        texto=str(row.get('Propuesta') or texto_contextual)
                    )
                except Exception:
                    fallos_no_recuperados += 1
                    print(
                        f'  [⚠ cross-eje] idx={idx} falló estricto y legado. '
                        'Se asume sin copias.'
                    )
                    matches = []
            else:
                fallos_no_recuperados += 1
                print(
                    f'  [⚠ cross-eje] idx={idx} falló estricto sin fallback. '
                    'Se asume sin copias.'
                )
                matches = []

        ejes_adicionales: list[str] = []

        for m in matches:
            eje_dest = str(m.get('eje', '')).strip()
            razon    = str(m.get('razon', '')).strip()

            try:
                confianza = int(m.get('confianza', 0))
            except (TypeError, ValueError):
                confianza = 0

            if eje_dest not in NOMBRES_EJES or eje_dest == eje_actual:
                continue

            ejes_adicionales.append(eje_dest)

            copia = row.to_dict()
            copia['Es_Copia']         = True
            copia['Eje_Origen_Copia'] = eje_actual
            copia['Razon_Cross_Eje']  = razon
            copia['Confianza_Cross']  = confianza
            copias_por_eje[eje_dest].append(copia)

        ejes_adicionales = _deduplicar_preservando_orden(ejes_adicionales)
        if ejes_adicionales:
            df_eje.at[idx, 'Ejes_Adicionales'] = ejes_adicionales

        time.sleep(PAUSA_ENTRE_LLAMADAS * 0.5)

    tasa_no_recuperados = (
        fallos_no_recuperados / evaluadas_cross
        if evaluadas_cross else 0.0
    )

    print(
        f'  → Cross-eje evaluadas: {evaluadas_cross} | '
        f'fallos estricto: {fallos_estrictos} | '
        f'fallback activado: {activaciones_fallback} | '
        f'fallos no recuperados: {fallos_no_recuperados} | '
        f'tasa no recuperados: {tasa_no_recuperados:.1%}'
    )

    if MODO_ESTRICTO_CROSS_EJE and fallos_no_recuperados > 0:
        raise RuntimeError(
            f'Se detectaron {fallos_no_recuperados} fallos no recuperados en cross-eje.'
        )

    if tasa_no_recuperados > TASA_MAX_FALLOS_CROSS_EJE:
        raise RuntimeError(
            f'Cross-eje demasiado inestable: {fallos_no_recuperados}/{evaluadas_cross} '
            f'fallos no recuperados ({tasa_no_recuperados:.1%}), por encima del máximo '
            f'permitido ({TASA_MAX_FALLOS_CROSS_EJE:.0%}).'
        )

    dfs_copias: dict[str, pd.DataFrame] = {
        eje: pd.DataFrame(lista)
        for eje, lista in copias_por_eje.items()
        if lista
    }

    for eje_dest, df_cop in dfs_copias.items():
        print(f'  → {len(df_cop)} copias salientes para «{eje_dest[:45]}»')

    return df_eje, dfs_copias


# ─────────────────────────────────────────────────────────────────────────────
# 4.  UTILIDADES WORD
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)


def _add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer  = section.footer
    fp      = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _limpiar_parrafo_compat(fp)

    run = fp.add_run('Página ')
    run.font.size = Pt(8)

    for instruccion in ['PAGE', 'NUMPAGES']:
        e = OxmlElement('w:fldChar')
        e.set(qn('w:fldCharType'), 'begin')
        run._r.append(e)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {instruccion} '
        run._r.append(instr)

        e2 = OxmlElement('w:fldChar')
        e2.set(qn('w:fldCharType'), 'end')
        run._r.append(e2)

        if instruccion == 'PAGE':
            run2 = fp.add_run(' de ')
            run2.font.size = Pt(8)
            run = run2

    pPr  = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'AAAAAA')
    pBdr.append(top)
    pPr.append(pBdr)


def _add_cover_page(doc: Document, eje: str, fecha: str, n_propuestas: int) -> None:
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(3)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, COLOR_AZUL_OSCURO)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(18)
    r = p.add_run('MECUN')
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Universidad Nacional de Colombia — Reforma Normativa')
    r2.font.size = Pt(12)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(COLOR_AZUL_MEDIO)

    doc.add_paragraph()
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c2 = tbl2.cell(0, 0)
    _set_cell_bg(c2, 'F2F2F2')
    pc = c2.paragraphs[0]
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_before = pc.paragraph_format.space_after = Pt(10)
    etiq = pc.add_run('EJE TEMÁTICO  ')
    etiq.bold = True
    etiq.font.size = Pt(9)
    etiq.font.color.rgb = RGBColor.from_string(COLOR_AZUL_MEDIO)
    rv = pc.add_run(eje)
    rv.bold = True
    rv.font.size = Pt(13)

    doc.add_paragraph()
    tbl3 = doc.add_table(rows=4, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ('Fecha de generación',      fecha),
        ('Propuestas procesadas',    str(n_propuestas)),
        ('Umbral de confianza IA',   f'{UMBRAL_CONFIANZA}%'),
        ('Umbral propagación links', f'{UMBRAL_PROPAGACION}/95'),
    ]):
        rk  = tbl3.rows[i].cells[0]
        rv2 = tbl3.rows[i].cells[1]
        _set_cell_bg(rk, COLOR_AZUL_OSCURO)
        r_k = rk.paragraphs[0].add_run(k)
        r_k.bold = True
        r_k.font.size = Pt(9)
        r_k.font.color.rgb = RGBColor(255, 255, 255)
        rv2.paragraphs[0].add_run(v).font.size = Pt(9)

    doc.add_paragraph()
    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (color, texto) in enumerate([
        (COLOR_VERDE,   '🟢 Alta confianza (≥70%)'),
        (COLOR_NARANJA, '🟡 Media confianza (40–69%)'),
        (COLOR_ROJO,    '🔴 Revisar (<40%)'),
    ]):
        c = tbl4.rows[0].cells[i]
        _set_cell_bg(c, COLOR_GRIS_CLARO)
        p_l = c.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_l = p_l.add_run(texto)
        r_l.font.size = Pt(8)
        r_l.font.color.rgb = RGBColor.from_string(color)

    doc.add_paragraph()
    p_ley = doc.add_paragraph()
    p_ley.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ley = p_ley.add_run('📎 = Relatoría adjunta   |   📎? = Relatoría propagada   |   🔀 = Copia de otro eje')
    r_ley.font.size = Pt(8)
    r_ley.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()


def _add_summary_table(doc: Document, resumen: list[dict]) -> None:
    doc.add_heading('Resumen por Categoría', level=1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Categoría', 'Propuestas', 'Grupos', 'Distribución']):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, COLOR_AZUL_OSCURO)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)

    total = sum(r['n_propuestas'] for r in resumen) or 1
    max_n = max((r['n_propuestas'] for r in resumen), default=1)

    for i, fila in enumerate(resumen):
        row = tbl.add_row()
        bg  = COLOR_GRIS_CLARO if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([
            fila['categoria'],
            str(fila['n_propuestas']),
            str(fila['n_grupos'])
        ]):
            _set_cell_bg(row.cells[ci], bg)
            row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

        c3 = row.cells[3]
        _set_cell_bg(c3, bg)
        bloques = int(round(fila['n_propuestas'] / max_n * 20)) if max_n else 0
        barra   = '█' * bloques + '░' * (20 - bloques)
        pct     = f'{fila["n_propuestas"]/total*100:.0f}%'
        r3 = c3.paragraphs[0].add_run(f'{barra} {pct}')
        r3.font.size = Pt(7)
        r3.font.name = 'Courier New'

    doc.add_paragraph()
    doc.add_page_break()


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    for tag, attrs in [
        ('w:color', {'w:val': '0563C1'}),
        ('w:u',     {'w:val': 'single'}),
        ('w:sz',    {'w:val': '16'}),
    ]:
        e = OxmlElement(tag)
        for k, v in attrs.items():
            e.set(qn(k), v)
        rPr.append(e)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_proposal_table(doc: Document, sintesis: str,
                         propuestas: list[dict], es_unica: bool = False) -> None:
    icono_color = COLOR_AZUL_OSCURO if not es_unica else '555555'
    icono_txt   = '💡 IDEA CENTRAL' if not es_unica else '📌 IDEA ÚNICA'

    tbl_h = doc.add_table(rows=1, cols=2)
    tbl_h.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1 = tbl_h.rows[0].cells
    _set_cell_bg(c0, icono_color)
    _set_cell_bg(c1, COLOR_AZUL_MEDIO if not es_unica else '444444')

    r0 = c0.paragraphs[0].add_run(icono_txt)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = RGBColor(255, 255, 255)

    r1 = c1.paragraphs[0].add_run(sintesis)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(255, 255, 255)

    tbl_v = doc.add_table(rows=1, cols=1)
    _set_cell_bg(tbl_v.rows[0].cells[0], 'FFF9E6')
    p_v = tbl_v.rows[0].cells[0].paragraphs[0]
    p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_v.add_run('⬜ VIABLE          ⬜ NO VIABLE          ⬜ REQUIERE AJUSTE').font.size = Pt(9)

    tbl_p = doc.add_table(rows=1, cols=4)
    tbl_p.style = 'Table Grid'
    for i, h in enumerate(['#', 'Claustro / Sede', 'Propuesta (texto original)', 'Relatoría']):
        c = tbl_p.rows[0].cells[i]
        _set_cell_bg(c, COLOR_AZUL_MEDIO)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(8)

    for j, prop in enumerate(propuestas):
        row = tbl_p.add_row()
        bg  = COLOR_GRIS_CLARO if j % 2 == 0 else 'FFFFFF'
        conf = int(prop.get('confianza', 0) or 0)

        c0 = row.cells[0]
        _set_cell_bg(c0, COLOR_LILA if prop.get('es_copia') else bg)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        id_txt = f'🔀 {prop["id"]}' if prop.get('es_copia') else str(prop['id'])
        r0 = p0.add_run(id_txt)
        r0.font.size = Pt(8)
        if prop.get('es_copia'):
            r0.font.color.rgb = RGBColor(255, 255, 255)

        c1 = row.cells[1]
        _set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.add_run(_safe_str(prop['claustro'])).font.size = Pt(7)
        r1b = p1.add_run(f"\n{_safe_str(prop['sede'])}")
        r1b.font.size = Pt(7)
        r1b.font.color.rgb = RGBColor(100, 100, 100)
        if prop.get('es_copia'):
            r1c = p1.add_run(f"\n🔀 Origen: {_safe_str(prop.get('eje_origen', ''))[:40]}")
            r1c.font.size = Pt(6)
            r1c.font.color.rgb = RGBColor.from_string(COLOR_LILA)

        c2 = row.cells[2]
        _set_cell_bg(c2, bg)
        c2.paragraphs[0].add_run(_safe_str(prop['texto'])).font.size = Pt(9)

        c3 = row.cells[3]
        _set_cell_bg(c3, bg)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_link   = prop.get('doc_link')
        doc_fuente = prop.get('doc_fuente', 'sin_doc')

        doc_link_str = _safe_str(doc_link)
        if doc_link_str and doc_fuente in ('original', 'propagado'):
            icono_link = '📎' if doc_fuente == 'original' else '📎?'
            p3.add_run(f'{icono_link} ').font.size = Pt(9)
            texto_link = doc_link_str[:30] + '…' if len(doc_link_str) > 30 else doc_link_str
            _add_hyperlink(p3, doc_link_str, texto_link)
            if doc_fuente == 'propagado':
                r3b = p3.add_run('\n(propagado)')
                r3b.font.size = Pt(6)
                r3b.font.color.rgb = RGBColor(120, 120, 120)
        else:
            r3x = p3.add_run('Sin relatoría')
            r3x.font.size = Pt(7)
            r3x.font.color.rgb = RGBColor(160, 160, 160)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# 5.  PREPROCESAMIENTO
# ─────────────────────────────────────────────────────────────────────────────

def _buscar_columna(columnas: list[str], *fragmentos: str,
                    excluir: tuple[str, ...] = ()) -> str | None:
    fl = [f.lower() for f in fragmentos]
    el = [e.lower() for e in excluir]
    for col in columnas:
        cl = col.lower()
        if all(f in cl for f in fl) and not any(e in cl for e in el):
            return col
    return None


def preprocesar_wide_a_long(ruta_csv: str) -> pd.DataFrame:
    print('Cargando CSV…')
    df_raw = leer_csv_robusto(ruta_csv)
    df_raw.columns = df_raw.columns.str.strip()
    cols = list(df_raw.columns)
    print(f'  → {len(df_raw)} filas × {len(df_raw.columns)} columnas')

    col_doc       = _buscar_columna(cols, 'Documento')
    col_claustro  = _buscar_columna(cols, 'Nombre del claustro')
    col_email     = _buscar_columna(cols, 'correo electrónico')
    col_tiempo    = _buscar_columna(cols, 'Marca temporal')
    col_sede      = _buscar_columna(cols, 'Sede de la reunión')
    col_estamento = _buscar_columna(cols, 'Estamentos que participan')

    print('Módulo A: Propagando links de relatoría…')
    df_raw = propagar_links_relatorias(
        df_raw, col_doc, col_claustro, col_email, col_tiempo
    )

    SLOTS_EJ  = ['EJ1', 'EJ2', 'EJ3', 'EJ4', 'EJ5']
    NUMS_PROP = ['1', '2', '3']
    filas, total_vacias = [], 0

    for _, fila in df_raw.iterrows():
        meta = {
            'Marca_Temporal'  : fila.get(col_tiempo,    'N/D') if col_tiempo else 'N/D',
            'Nombre_Claustro' : fila.get(col_claustro,  'N/D') if col_claustro else 'N/D',
            'Email'           : fila.get(col_email,     'N/D') if col_email else 'N/D',
            'Sede'            : fila.get(col_sede,      'N/D') if col_sede else 'N/D',
            'Estamentos'      : fila.get(col_estamento, 'N/D') if col_estamento else 'N/D',
            'Documento_Link'  : fila.get('Documento_Propagado', None),
            'Doc_Score'       : fila.get('Doc_Score', 0),
            'Doc_Fuente'      : fila.get('Doc_Fuente', 'sin_doc'),
            'Doc_Razones'     : fila.get('Doc_Razones', ''),
        }

        for slot in SLOTS_EJ:
            eje_del_slot = EJES_VALIDOS.get(slot, 'N/D')

            for num in NUMS_PROP:
                col_prop   = _buscar_columna(cols, f'[{slot}]', f'Propuesta {num}',
                                              excluir=('diagnóstico', 'título'))
                col_titulo = _buscar_columna(cols, f'[{slot}]', 'Título', f'{num}')
                col_diag   = _buscar_columna(cols, f'[{slot}]', 'Diagnóstico', f'{num}')

                if col_prop is None:
                    continue

                val = fila.get(col_prop, None)
                if pd.isna(val) or str(val).strip() in ('', 'nan', 'NaN'):
                    total_vacias += 1
                    continue

                propuesta   = str(val).strip()
                titulo      = str(fila.get(col_titulo, 'Sin título')).strip() if col_titulo else 'Sin título'
                diagnostico = str(fila.get(col_diag,   'N/D')).strip() if col_diag else 'N/D'

                if titulo in ('nan', '', 'NaN'):
                    titulo = 'Sin título'
                if diagnostico in ('nan', '', 'NaN'):
                    diagnostico = 'N/D'

                filas.append({
                    **meta,
                    'Eje_Tematico'         : eje_del_slot,
                    'Slot_EJ'              : slot,
                    'Num_Propuesta'        : num,
                    'Titulo'               : titulo,
                    'Diagnostico'          : diagnostico,
                    'Propuesta'            : propuesta,
                    'Texto_Contextualizado': (
                        f'TÍTULO: {titulo} | '
                        f'DIAGNÓSTICO: {diagnostico} | '
                        f'PROPUESTA: {propuesta}'
                    ),
                    'Es_Copia'             : False,
                    'Eje_Origen_Copia'     : '',
                    'Ejes_Adicionales'     : [],
                })

    df_largo = pd.DataFrame(filas)
    print(f'\n✅ Preprocesamiento: {len(df_largo)} propuestas válidas '
          f'({total_vacias} celdas vacías ignoradas)')
    for eje, cnt in df_largo['Eje_Tematico'].value_counts().items():
        print(f'   {cnt:>4}  {eje}')
    return df_largo


def _asegurar_lista_segura(x):
    if isinstance(x, list):
        return x
    if isinstance(x, tuple):
        return list(x)
    if isinstance(x, set):
        return list(x)

    if x is None:
        return []

    try:
        if pd.isna(x):
            return []
    except Exception:
        pass

    if isinstance(x, str):
        s = x.strip()
        if s in ('', 'nan', 'NaN', 'None', 'null'):
            return []
        return [s]

    return []


# ─────────────────────────────────────────────────────────────────────────────
# 6.  PIPELINE MAESTRO — TODOS LOS EJES
# ─────────────────────────────────────────────────────────────────────────────

def procesar_eje_completo(df_eje: pd.DataFrame, eje_actual: str) -> None:
    global EJE_ACTUAL, CATEGORIAS_BASE, CARPETA_SALIDA

    EJE_ACTUAL      = eje_actual
    CATEGORIAS_BASE = DICCIONARIO_CATEGORIAS.get(EJE_ACTUAL, [])
    CARPETA_SALIDA  = crear_carpeta_salida(EJE_ACTUAL)
    df_eje          = df_eje.copy()

    columnas_texto = [
        'Propuesta', 'Texto_Contextualizado', 'Titulo', 'Diagnostico',
        'Nombre_Claustro', 'Sede', 'Documento_Link', 'Doc_Fuente',
        'Eje_Origen_Copia'
    ]
    for col in columnas_texto:
        if col not in df_eje.columns:
            df_eje[col] = ''
        df_eje[col] = df_eje[col].apply(_safe_str)

    if 'Confianza' not in df_eje.columns:
        df_eje['Confianza'] = 0
    else:
        df_eje['Confianza'] = pd.to_numeric(df_eje['Confianza'], errors='coerce').fillna(0).astype(int)

    if 'Es_Copia' not in df_eje.columns:
        df_eje['Es_Copia'] = False
    else:
        df_eje['Es_Copia'] = df_eje['Es_Copia'].fillna(False).astype(bool)

    if 'Ejes_Adicionales' not in df_eje.columns:
        df_eje['Ejes_Adicionales'] = [[] for _ in range(len(df_eje))]
    else:
        df_eje['Ejes_Adicionales'] = df_eje['Ejes_Adicionales'].apply(_asegurar_lista_segura)

    print(f'\nEje: «{EJE_ACTUAL}»  |  Propuestas: {len(df_eje)}\n')

    # ─── Parámetros metodológicos locales ────────────────────────────────────
    FRACCION_MUESTRA_DESC        = 0.35
    TAMANO_MIN_MUESTRA_DESC      = 40
    TAMANO_MAX_MUESTRA_DESC      = 140
    VALIDAR_EN_SEGUNDA_MUESTRA   = True
    APOYO_MINIMO_CATEGORIA_NUEVA = 3
    MIN_PROPUESTAS_POR_CATEGORIA = 4

    # ─── Helpers locales de descubrimiento ───────────────────────────────────
    def calcular_muestra_descubrimiento_rigurosa(total: int) -> int:
        if total <= 0:
            return 0
        n_slovin = calcular_muestra_estadistica(total, MARGEN_ERROR_SLOVIN)
        n_frac   = int(round(total * FRACCION_MUESTRA_DESC))
        n        = max(n_slovin, n_frac, TAMANO_MIN_MUESTRA_DESC)
        return min(n, total, TAMANO_MAX_MUESTRA_DESC)

    def construir_muestra_descubrimiento(df: pd.DataFrame,
                                          n: int,
                                          random_state: int = 42) -> list[str]:
        if df.empty or n <= 0:
            return []

        base = df.copy()
        base = base.sample(frac=1, random_state=random_state).reset_index(drop=True)

        grupos_cols = [c for c in ['Sede', 'Nombre_Claustro'] if c in base.columns]
        if grupos_cols:
            muestras    = []
            grupos      = list(base.groupby(grupos_cols, dropna=False))
            cupo_base   = max(1, n // max(len(grupos), 1))
            usados      = set()
            for _, g in grupos:
                g_take = g.head(min(len(g), cupo_base))
                muestras.append(g_take)
                usados.update(g_take.index.tolist())

            restante = n - sum(len(x) for x in muestras)
            if restante > 0:
                faltantes = base.loc[~base.index.isin(usados)]
                if not faltantes.empty:
                    muestras.append(faltantes.head(restante))

            muestra_df = pd.concat(muestras, ignore_index=True) if muestras else base.head(n)
            muestra_df = muestra_df.head(n)
        else:
            muestra_df = base.head(n)

        return (
            muestra_df['Texto_Contextualizado']
            .dropna()
            .astype(str)
            .tolist()
        )

    def contar_apoyo_categoria_en_textos(categoria: str,
                                          textos: list[str],
                                          model: str = MODEL_FLASH) -> int:
        if not textos:
            return 0

        prompt = (
            'Evalúa cuántos textos muestran claramente el patrón temático descrito por esta categoría.\n'
            f'Categoría candidata: {categoria}\n\n'
            'Cuenta SOLO apoyos claros y directos.\n'
            'Responde SOLO con un entero entre 0 y el número total de textos.\n\n'
            'TEXTOS:\n' + '\n---\n'.join(textos)
        )
        respuesta = llamar_api(prompt, model=model)
        if not respuesta:
            return 0

        m = re.search(r'\b(\d{1,3})\b', respuesta)
        if not m:
            return 0

        try:
            return max(0, min(int(m.group(1)), len(textos)))
        except Exception:
            return 0

    def _filtrar_categorias_dinamicas_riguroso(categorias_raw: list[str],
                                                categorias_base: list[str]) -> list[str]:
        resultado: list[str] = []
        universo_existente = [c for c in categorias_base if 'Casos Aislados' not in c]

        for cat in categorias_raw:
            c = _normalizar_etiqueta_categoria(cat)

            if len(c) < LONGITUD_MIN_CATEGORIA:
                continue

            c_low = c.lower()
            if c_low in {'ninguna', 'otros', 'otras', 'varios', 'miscelánea', 'miscelanea'}:
                continue

            if any(_parecido(c, base) >= SIMILITUD_MAX_CATEGORIAS for base in universo_existente):
                continue

            if any(_parecido(c, ya) >= SIMILITUD_MAX_CATEGORIAS for ya in resultado):
                continue

            resultado.append(c)

        return resultado

    # ─── FASE 1 — DESCUBRIMIENTO DE CATEGORÍAS ───────────────────────────────
    tamano_muestra = calcular_muestra_descubrimiento_rigurosa(len(df_eje))
    print(f'\nFase 1: Descubrimiento → muestra rigurosa {tamano_muestra}/{len(df_eje)}')

    df_no_copia = df_eje[~df_eje['Es_Copia']].copy()
    if df_no_copia.empty:
        raise RuntimeError('No hay propuestas originales para construir muestra de descubrimiento.')

    df_desc  = df_no_copia.sample(frac=1, random_state=42).reset_index(drop=True)
    tamano_muestra     = min(tamano_muestra, len(df_desc))
    tamano_validacion  = max(0, min(len(df_desc) - tamano_muestra, max(20, tamano_muestra // 2)))

    df_muestra = df_desc.iloc[:tamano_muestra].copy()
    df_valid   = df_desc.iloc[tamano_muestra:tamano_muestra + tamano_validacion].copy()

    textos_muestra    = construir_muestra_descubrimiento(df_muestra, len(df_muestra), random_state=42)
    textos_validacion = (
        df_valid['Texto_Contextualizado'].dropna().astype(str).tolist()
        if not df_valid.empty else []
    )

    categorias_dinamicas: list[str] = []

    if USAR_DESCUBRIMIENTO_DINAMICO:
        prompt_desc = (
            'Eres un analista evaluando propuestas universitarias (MECUN).\n'
            'Categorías Base ya existentes (NO las repitas ni reformules):\n'
            + '\n'.join(f'- {c}' for c in CATEGORIAS_BASE)
            + '\n\nIdentifica SOLO patrones nuevos, claramente distintos de las categorías base.\n'
            'Restricciones:\n'
            '- no inventes sinónimos de categorías base\n'
            '- no subdividas artificialmente una categoría base ya suficiente\n'
            '- si una propuesta cabe razonablemente en una categoría base, NO propongas una nueva\n'
            '- no uses categorías vagas como "otros", "varios", "general"\n'
            '- cada categoría nueva debe agrupar varias propuestas, no casos aislados\n'
            '- responde SOLO con lista enumerada o NINGUNA\n\n'
            'MUESTRA:\n'
            + '\n---\n'.join(textos_muestra)
        )

        respuesta_desc = llamar_api(prompt_desc, model=MODEL_FLASH)

        categorias_raw = []
        if respuesta_desc and 'NINGUNA' not in respuesta_desc.upper() and len(respuesta_desc) > 5:
            for linea in respuesta_desc.split('\n'):
                txt = linea.lstrip('0123456789.-* ').strip()
                if txt:
                    categorias_raw.append(txt)

        categorias_dinamicas = _filtrar_categorias_dinamicas_riguroso(
            categorias_raw,
            CATEGORIAS_BASE
        )

        if categorias_dinamicas and VALIDAR_EN_SEGUNDA_MUESTRA and textos_validacion:
            categorias_validadas: list[str] = []
            for cat in categorias_dinamicas:
                apoyo = contar_apoyo_categoria_en_textos(cat, textos_validacion)
                if apoyo >= APOYO_MINIMO_CATEGORIA_NUEVA:
                    categorias_validadas.append(cat)
            categorias_dinamicas = categorias_validadas

        if categorias_dinamicas:
            print(f'  → {len(categorias_dinamicas)} categorías nuevas aceptadas tras filtrado y validación')
        else:
            print('  → Categorías base suficientes')
    else:
        print('  → Descubrimiento dinámico desactivado; se usarán solo categorías base.')

    categorias_totales = (
        [c for c in CATEGORIAS_BASE if 'Casos Aislados' not in c]
        + categorias_dinamicas
        + ['Casos Aislados / Otros']
    )

    if len(categorias_totales) < 2:
        raise RuntimeError('El conjunto de categorías quedó inválido.')

    texto_cats = '\n'.join(f'{i+1}. {c}' for i, c in enumerate(categorias_totales))
    time.sleep(PAUSA_ENTRE_LLAMADAS)

    # ─── FASE 2 — CLASIFICACIÓN ──────────────────────────────────────────────
    print(f'\nFase 2: Clasificación por lotes de {TAMANO_LOTE_CLASIF}…')

    PROMPT_LOTE = (
        f'Clasifica CADA propuesta en UNA de estas categorías numeradas:\n{texto_cats}\n\n'
        'INSTRUCCIONES CRÍTICAS:\n'
        '- Responde SOLO con JSON válido.\n'
        '- Usa el número de la categoría, NO el nombre.\n'
        '- El campo debe llamarse exactamente "categoria_id".\n'
        f'- "categoria_id" debe ser un entero entre 1 y {len(categorias_totales)}.\n'
        '- No agregues texto antes ni después del JSON.\n\n'
        'Formato exacto:\n'
        '[{"id": 1, "categoria_id": 3, "confianza": 85}, ...]\n\n'
        'Propuestas:\n{propuestas_bloque}'
    )

    def clasificar_lote_con_rescate(lote_ids: list[int]) -> list[dict]:
        def _intentar(ids_locales: list[int], contexto_extra: str) -> tuple[list[dict], bool]:
            bloque_props_local = '\n'.join(
                f'[{id_}] {df_eje.iloc[id_-1]["Texto_Contextualizado"]}'
                for id_ in ids_locales
            )

            respuesta_local = llamar_api(
                PROMPT_LOTE.replace('{propuestas_bloque}', bloque_props_local),
                model=MODEL_FLASH
            )

            items_local, exito_local = _parsear_json_respuesta(
                respuesta_local,
                ids_locales,
                contexto=f'Fase2/{eje_actual[:20]}/{contexto_extra}',
                validador=_validar_items_clasificacion_por_id
            )

            if not exito_local:
                return [], False

            ids_vistos_local = set()
            items_saneados   = []

            for item in items_local:
                if not isinstance(item, dict):
                    return [], False

                try:
                    rid          = int(item.get('id'))
                    categoria_id = int(item.get('categoria_id'))
                    conf         = int(item.get('confianza', 0))
                except (TypeError, ValueError):
                    return [], False

                if rid not in ids_locales:
                    return [], False

                if not (1 <= categoria_id <= len(categorias_totales)):
                    return [], False

                if rid in ids_vistos_local:
                    return [], False

                ids_vistos_local.add(rid)
                items_saneados.append({
                    'id': rid,
                    'categoria_id': categoria_id,
                    'confianza': conf
                })

            faltantes = [id_ for id_ in ids_locales if id_ not in ids_vistos_local]
            if faltantes:
                return [], False

            return items_saneados, True

        items_ok, ok = _intentar(
            lote_ids,
            f'lote {lote_ids[0]}-{lote_ids[-1]} / intento_1'
        )
        if ok:
            return items_ok

        time.sleep(PAUSA_ENTRE_LLAMADAS)

        items_ok, ok = _intentar(
            lote_ids,
            f'lote {lote_ids[0]}-{lote_ids[-1]} / intento_2'
        )
        if ok:
            print(f'  [rescate] lote {lote_ids[0]}-{lote_ids[-1]} recuperado en segundo intento')
            return items_ok

        time.sleep(PAUSA_ENTRE_LLAMADAS)

        print(f'  [rescate] lote {lote_ids[0]}-{lote_ids[-1]} cayó a clasificación individual')
        acumulado: list[dict] = []

        for rid in lote_ids:
            item_1, ok_1 = _intentar([rid], f'id {rid} / individual_1')

            if not ok_1:
                time.sleep(PAUSA_ENTRE_LLAMADAS)
                item_1, ok_1 = _intentar([rid], f'id {rid} / individual_2')

            if not ok_1 or not item_1:
                print(
                    f'  [seguridad] id={rid} no pudo clasificarse automáticamente. '
                    f'Se envía a "Casos Aislados / Otros".'
                )
                acumulado.append({
                    'id': rid,
                    'categoria_id': len(categorias_totales),
                    'confianza': 0
                })
            else:
                acumulado.extend(item_1)

            time.sleep(PAUSA_ENTRE_LLAMADAS * 0.5)

        ids_acum = sorted(int(x['id']) for x in acumulado)
        if ids_acum != sorted(lote_ids):
            raise RuntimeError(
                f'Fallback individual incompleto en lote {lote_ids[0]}-{lote_ids[-1]}. '
                f'Esperados={lote_ids} | Obtenidos={ids_acum}'
            )

        return acumulado

    todos_ids       = list(range(1, len(df_eje) + 1))
    etiquetas_dict: dict[int, str]  = {}
    confianzas_dict: dict[int, int] = {}
    lotes = [todos_ids[i:i+TAMANO_LOTE_CLASIF] for i in range(0, len(todos_ids), TAMANO_LOTE_CLASIF)]

    for lote_ids in tqdm(lotes, desc=f'Clasificando {eje_actual[:35]}', unit='lote'):
        items = clasificar_lote_con_rescate(lote_ids)

        ids_observados = set()

        for item in items:
            rid          = int(item['id'])
            categoria_id = int(item['categoria_id'])
            conf         = int(item.get('confianza', 0))

            cat = categorias_totales[categoria_id - 1]

            if rid in ids_observados:
                raise RuntimeError(f'ID duplicado en respuesta final de Fase 2: {rid}')

            ids_observados.add(rid)
            etiquetas_dict[rid]  = cat
            confianzas_dict[rid] = conf

        faltantes_lote = [id_ for id_ in lote_ids if id_ not in ids_observados]
        if faltantes_lote:
            raise RuntimeError(
                f'Fase 2 dejó propuestas sin clasificar en lote {lote_ids[0]}-{lote_ids[-1]}: {faltantes_lote}'
            )

        time.sleep(PAUSA_ENTRE_LLAMADAS)

    df_eje = df_eje.copy()
    df_eje['Categoria_Final'] = [etiquetas_dict.get(i, 'Casos Aislados / Otros') for i in todos_ids]
    df_eje['Confianza']       = [confianzas_dict.get(i, 0) for i in todos_ids]

    if len(df_eje) != len(todos_ids):
        raise RuntimeError(
            f'Desajuste crítico: df_eje tiene {len(df_eje)} filas pero todos_ids tiene {len(todos_ids)}.'
        )

    if df_eje['Categoria_Final'].isna().any():
        faltantes = df_eje[df_eje['Categoria_Final'].isna()].index.tolist()
        raise RuntimeError(f'Se detectaron propuestas sin categoría asignada: {faltantes[:20]}')

    conteo_categorias  = df_eje['Categoria_Final'].value_counts()
    categorias_raquiticas = [
        cat for cat, n in conteo_categorias.items()
        if cat != 'Casos Aislados / Otros' and n < MIN_PROPUESTAS_POR_CATEGORIA
    ]
    if categorias_raquiticas:
        df_eje.loc[
            df_eje['Categoria_Final'].isin(categorias_raquiticas),
            'Categoria_Final'
        ] = 'Casos Aislados / Otros'
        print(
            f'  → {len(categorias_raquiticas)} categorías con baja densidad '
            'se reagruparon en "Casos Aislados / Otros"'
        )

    # ─── FASE 3 — AGRUPAMIENTO Y WORD EJECUTIVO ──────────────────────────────
    print('\nFase 3: Agrupando y generando Word ejecutivo…')

    categorias_unicas = sorted([
        c for c in df_eje['Categoria_Final'].dropna().unique()
        if str(c).strip() not in ('', 'Ninguna')
    ])

    datos_excel    = []
    resumen_cats   = []
    grupos_por_cat = {}

    PROMPT_AGRUPAR = (
        'Agrupa estas {n} propuestas de la categoría "{cat}" por IDEA CENTRAL compartida.\n'
        'REGLA 1: responde SOLO un arreglo JSON válido.\n'
        'REGLA 2: cada elemento debe ser un objeto con EXACTAMENTE estas llaves: '
        '"sintesis" (string) e "ids" (lista).\n'
        'REGLA 3: no incluyas texto adicional, markdown, comentarios ni explicación.\n'
        'REGLA 4: usa únicamente ids presentes en el lote.\n'
        'REGLA 5: cada id debe aparecer una sola vez.\n'
        'REGLA 6: si una propuesta no encaja con otras, déjala como grupo de un solo id.\n'
        'REGLA 7: si no puedes agrupar con certeza, devuelve []\n'
        'REGLA 8: "sintesis" debe ser una frase breve, máximo 15 palabras.\n'
        'Formato exacto esperado:\n'
        '[{{"sintesis": "Frase breve", "ids": [1,4,7]}}]\n\n'
        'Propuestas (Lote {lote_num}):\n{bloque}'
    )

    for idx_cat, categoria in enumerate(tqdm(categorias_unicas, desc=f'Agrupando {eje_actual[:35]}')):
        df_grupo = df_eje[df_eje['Categoria_Final'] == categoria].reset_index(drop=True)
        mapeo = {
            str(i + 1): {
                'id'         : i + 1,
                'texto'      : row['Propuesta'],
                'titulo'     : row['Titulo'],
                'diagnostico': row['Diagnostico'],
                'claustro'   : row['Nombre_Claustro'],
                'sede'       : row['Sede'],
                'confianza'  : row['Confianza'],
                'doc_link'   : row.get('Documento_Link'),
                'doc_fuente' : row.get('Doc_Fuente', 'sin_doc'),
                'doc_score'  : row.get('Doc_Score', 0),
                'es_copia'   : bool(row.get('Es_Copia', False)),
                'eje_origen' : row.get('Eje_Origen_Copia', ''),
            }
            for i, (_, row) in enumerate(df_grupo.iterrows())
        }

        if not mapeo:
            raise RuntimeError(f'Categoría "{categoria}" sin propuestas mapeables.')

        items          = list(mapeo.items())
        bloques_grupos = []
        grupos_fallidos = 0

        for i in range(0, len(items), TAMANO_LOTE_AGRUPAR):
            lote       = items[i:i + TAMANO_LOTE_AGRUPAR]
            lote_ids   = [k for k, _ in lote]
            bloque_txt = '\n'.join(f'[{k}] {v["texto"]}' for k, v in lote)

            resp_ag = llamar_api(
                PROMPT_AGRUPAR.format(
                    n=len(lote),
                    cat=categoria,
                    lote_num=i // TAMANO_LOTE_AGRUPAR + 1,
                    bloque=bloque_txt
                ),
                model=MODEL_FLASH
            )

            grupos_lote, exito_ag = _parsear_json_respuesta(
                resp_ag,
                lote_ids,
                contexto=f'Fase3/{categoria[:30]}/lote{i//TAMANO_LOTE_AGRUPAR+1}',
                validador=_validar_items_agrupacion
            )
            if not exito_ag:
                grupos_fallidos += 1
                if MODO_ESTRICTO_NUCLEO:
                    raise RuntimeError(
                        f'Parse fallido en Fase 3 para categoría "{categoria}", lote {i//TAMANO_LOTE_AGRUPAR+1}.'
                    )

            ids_lote_vistos      = set()
            grupos_lote_saneados = []

            for grupo in grupos_lote:
                if not isinstance(grupo, dict):
                    raise RuntimeError(f'Elemento no dict en Fase 3 categoría "{categoria}".')

                sintesis_raw = grupo.get('sintesis', '')
                ids_raw      = grupo.get('ids', [])

                if not isinstance(sintesis_raw, str):
                    raise RuntimeError(f'"sintesis" inválida en Fase 3 categoría "{categoria}".')

                ids_norm = _normalizar_ids_validos(ids_raw, dict(lote))
                if not ids_norm:
                    raise RuntimeError(f'Grupo vacío o sin ids válidos en Fase 3 categoría "{categoria}".')

                repetidos = [x for x in ids_norm if x in ids_lote_vistos]

                if repetidos:
                    print(
                        f'  [WARN] IDs repetidos entre grupos del mismo lote en Fase 3 '
                        f'categoría "{categoria}": {repetidos}. '
                        'Se conservará solo la primera aparición.'
                    )
                    ids_norm = [x for x in ids_norm if x not in ids_lote_vistos]

                if not ids_norm:
                    continue

                ids_lote_vistos.update(ids_norm)
                grupos_lote_saneados.append({
                    'sintesis': sintesis_raw.strip(),
                    'ids': ids_norm
                })

            bloques_grupos.extend(grupos_lote_saneados)
            time.sleep(PAUSA_ENTRE_LLAMADAS)

        if grupos_fallidos:
            raise RuntimeError(f'Categoría "{categoria}" tuvo {grupos_fallidos} lotes fallidos en Fase 3.')

        ids_cubiertos  = set()
        grupos_finales = []

        for grupo in bloques_grupos:
            sin   = str(grupo.get('sintesis', '')).strip()
            ids_v = _normalizar_ids_validos(grupo.get('ids', []), mapeo)

            if not sin:
                raise RuntimeError(f'Grupo sin síntesis en categoría "{categoria}".')

            ids_nuevos    = [x for x in ids_v if x not in ids_cubiertos]
            ids_repetidos = [x for x in ids_v if x in ids_cubiertos]

            if ids_repetidos:
                raise RuntimeError(
                    f'ID repetido entre grupos finales en categoría "{categoria}": {ids_repetidos}'
                )

            if not ids_nuevos:
                raise RuntimeError(f'Grupo sin ids nuevos en categoría "{categoria}".')

            grupos_finales.append({'sintesis': sin, 'ids': ids_nuevos})
            ids_cubiertos.update(ids_nuevos)

        for k in mapeo:
            if k not in ids_cubiertos:
                grupos_finales.append({
                    'sintesis': f'[Idea única] {mapeo[k]["titulo"]}',
                    'ids': [k],
                    '_unica': True,
                })

        _validar_cobertura_categoria(categoria, mapeo, grupos_finales)

        grupos_por_cat[categoria] = grupos_finales
        resumen_cats.append({
            'categoria'   : categoria,
            'n_propuestas': len(df_grupo),
            'n_grupos'    : len(grupos_finales),
        })

        for grupo in grupos_finales:
            for id_p in grupo['ids']:
                d = mapeo[id_p]
                datos_excel.append({
                    'Eje Temático'                     : EJE_ACTUAL,
                    'Categoría Asignada'               : categoria,
                    'Síntesis del Grupo (Idea Central)': grupo['sintesis'],
                    'Claustro'                         : d['claustro'],
                    'Sede'                             : d['sede'],
                    'Título Original'                  : d['titulo'],
                    'Diagnóstico'                      : d['diagnostico'],
                    'Texto Propuesta Original'         : d['texto'],
                    'Confianza Clasificación (%)'      : d['confianza'],
                    'Link Relatoría'                   : d['doc_link'] or '',
                    'Fuente Link'                      : d['doc_fuente'],
                    'Score Propagación'                : d['doc_score'],
                    'Es Copia Cross-Eje'               : d['es_copia'],
                    'Eje Origen (si copia)'            : d['eje_origen'],
                })

    if not resumen_cats:
        raise RuntimeError('No hay resumen de categorías; se aborta antes de exportar.')

    if not datos_excel:
        raise RuntimeError('No hay datos tabulares finales; se aborta antes de exportar.')

    if len(datos_excel) != len(df_eje):
        raise RuntimeError(
            f'Pérdida detectada antes de exportar: datos_excel={len(datos_excel)} '
            f'vs propuestas procesadas={len(df_eje)}'
        )

    doc = Document()
    _add_page_number_footer(doc)
    _add_cover_page(doc, EJE_ACTUAL, FECHA_EJECUCION, len(df_eje))
    _add_summary_table(doc, resumen_cats)

    for idx_cat, categoria in enumerate(categorias_unicas):
        df_grupo = df_eje[df_eje['Categoria_Final'] == categoria]
        grupos   = grupos_por_cat.get(categoria, [])

        if not grupos:
            raise RuntimeError(f'Categoría "{categoria}" no tiene grupos finales válidos.')

        if idx_cat > 0:
            doc.add_page_break()

        ph = doc.add_heading(categoria, level=1)
        ph.runs[0].font.color.rgb = RGBColor.from_string(COLOR_ROJO)

        p_meta = doc.add_paragraph()
        r_m1   = p_meta.add_run(f'{len(df_grupo)} propuestas  |  {len(grupos)} grupos')
        r_m1.font.size = Pt(9)
        r_m1.font.color.rgb = RGBColor.from_string(COLOR_AZUL_MEDIO)

        n_copias = int(df_grupo['Es_Copia'].sum()) if 'Es_Copia' in df_grupo else 0
        if n_copias:
            r_m2 = p_meta.add_run(f'  |  🔀 {n_copias} copias cross-eje')
            r_m2.font.size = Pt(9)
            r_m2.font.color.rgb = RGBColor.from_string(COLOR_LILA)

        doc.add_paragraph()

        mapeo_cat = {
            str(i + 1): {
                'id'        : i + 1,
                'texto'     : row['Propuesta'],
                'claustro'  : row['Nombre_Claustro'],
                'sede'      : row['Sede'],
                'confianza' : row['Confianza'],
                'doc_link'  : row.get('Documento_Link'),
                'doc_fuente': row.get('Doc_Fuente', 'sin_doc'),
                'doc_score' : row.get('Doc_Score', 0),
                'es_copia'  : bool(row.get('Es_Copia', False)),
                'eje_origen': row.get('Eje_Origen_Copia', ''),
            }
            for i, (_, row) in enumerate(df_grupo.reset_index(drop=True).iterrows())
        }

        for grupo in grupos:
            es_unica = bool(grupo.get('_unica', len(grupo['ids']) == 1))
            props    = [mapeo_cat[x] for x in grupo['ids'] if x in mapeo_cat]
            if not props:
                raise RuntimeError(f'Grupo sin props renderizables en categoría "{categoria}".')
            _add_proposal_table(doc, grupo['sintesis'], props, es_unica=es_unica)

    nombre_base = re.sub(r'[^a-z0-9]', '_', EJE_ACTUAL.lower())[:18].strip('_')
    nombre_word = CARPETA_SALIDA / f'Reporte_Ejecutivo_{nombre_base}_{TS}.docx'
    nombre_csv  = CARPETA_SALIDA / f'Reporte_Tabular_{nombre_base}_{TS}.csv'

    doc.save(nombre_word)
    print(f'\n  → Word: {nombre_word}')
    pd.DataFrame(datos_excel).to_csv(nombre_csv, index=False, sep=';', encoding='utf-8-sig')
    print(f'  → CSV:  {nombre_csv}')

    # ─── FASE 4 — CONSENSOS Y DISENSOS ───────────────────────────────────────
    print(f'\nFase 4: Mapa de Consensos y Disensos…')
    doc_sem = Document()
    _add_page_number_footer(doc_sem)
    _add_cover_page(
        doc_sem,
        f'CONSENSOS Y DISENSOS — {EJE_ACTUAL}',
        FECHA_EJECUCION,
        len(df_eje)
    )

    PROMPT_JUEZ = (
        'Analista experto en reforma normativa UNAL (MECUN).\n'
        'Eje: {eje} | Categoría: {cat} | Propuestas: {n}\n\n'
        'Lee íntegros los textos. Identifica SOLO los consensos/disensos reales.\n\n'
        '🟢 CONSENSOS MAYORITARIOS:\n- [Idea] - [Por qué es consenso].\n\n'
        '🟡 PUNTOS DE NEGOCIACIÓN:\n- [Punto] - [Diferentes posturas].\n\n'
        '🔴 DISENSOS IRRECONCILIABLES:\n- [Choque] - [Posturas opuestas]. '
        'O: "No se detectaron disensos irreconciliables."\n\nPROPUESTAS:\n{propuestas}'
    )

    for idx_cat, categoria in enumerate(tqdm(categorias_unicas, desc=f'Semáforo {eje_actual[:30]}')):
        df_grupo = df_eje[df_eje['Categoria_Final'] == categoria]
        if len(df_grupo) < 2:
            if idx_cat > 0:
                doc_sem.add_page_break()
            doc_sem.add_heading(f'Categoría: {categoria}', level=2)
            doc_sem.add_paragraph('Solo 1 propuesta — análisis omitido.').italic = True
            continue

        bloque_props = '\n\n'.join(
            f'[{i+1}] {row["Propuesta"]}'
            for i, (_, row) in enumerate(df_grupo.iterrows())
        )
        resp_j = llamar_api(
            PROMPT_JUEZ.format(
                eje=EJE_ACTUAL,
                cat=categoria,
                n=len(df_grupo),
                propuestas=bloque_props
            ),
            model=MODEL_PRO
        )

        if not resp_j:
            raise RuntimeError(f'La Fase 4 no devolvió contenido para categoría "{categoria}".')

        if idx_cat > 0:
            doc_sem.add_page_break()

        ph = doc_sem.add_heading(f'Categoría: {categoria}', level=2)
        ph.runs[0].font.color.rgb = RGBColor.from_string(COLOR_ROJO)
        doc_sem.add_paragraph(f'Análisis de {len(df_grupo)} propuestas íntegras.').italic = True

        for linea in resp_j.split('\n'):
            lc = linea.strip()
            if not lc:
                continue
            p = doc_sem.add_paragraph()
            if lc.startswith(('🟢', '🟡', '🔴')):
                r = p.add_run(lc)
                r.bold = True
                r.font.size = Pt(11)
                if lc.startswith('🟢'):
                    r.font.color.rgb = RGBColor.from_string(COLOR_VERDE)
                elif lc.startswith('🟡'):
                    r.font.color.rgb = RGBColor.from_string(COLOR_NARANJA)
                else:
                    r.font.color.rgb = RGBColor.from_string(COLOR_ROJO)
            elif lc.startswith('-'):
                p.style = 'List Bullet'
                partes  = lc.split('-', 2)
                if len(partes) > 2:
                    p.add_run('- ' + partes[1].strip() + ' - ').bold = True
                    p.add_run(partes[2].strip()).font.size = Pt(9)
                else:
                    p.add_run(lc).font.size = Pt(9)
            else:
                p.add_run(lc).font.size = Pt(9)

        time.sleep(PAUSA_ENTRE_LLAMADAS)

    nombre_sem = CARPETA_SALIDA / f'Semaforo_{nombre_base}_{TS}.docx'
    doc_sem.save(nombre_sem)
    print(f'  → Semáforo: {nombre_sem}')

    dataset_nombre = Path(DATA_SOURCE_PATH).stem
    print(f"""
🎉 PIPELINE v4.6 COMPLETADO — {dataset_nombre}
   Eje: {EJE_ACTUAL}
   Archivos en: {CARPETA_SALIDA}/
   • {nombre_word.name}
   • {nombre_csv.name}
   • {nombre_sem.name}
""")


# ─────────────────────────────────────────────────────────────────────────────
# 7.  CHECKPOINTS
# ─────────────────────────────────────────────────────────────────────────────

def guardar_checkpoint_pass1(df_largo, originales_por_eje, copias_entrantes_por_eje):
    with open(PASS1_CHECKPOINT, 'wb') as f:
        pickle.dump({
            'df_largo': df_largo,
            'originales_por_eje': originales_por_eje,
            'copias_entrantes_por_eje': copias_entrantes_por_eje,
        }, f)
    print(f'  💾 Checkpoint PASS 1 guardado en: {PASS1_CHECKPOINT}')


def cargar_checkpoint_pass1():
    if FORZAR_RECALCULO_PASS1:
        print('  ♻️ FORZAR_RECALCULO_PASS1=True → se ignorará el checkpoint.')
        return None

    if not PASS1_CHECKPOINT.exists():
        return None

    with open(PASS1_CHECKPOINT, 'rb') as f:
        data = pickle.load(f)

    print(f'  📦 Checkpoint PASS 1 cargado desde: {PASS1_CHECKPOINT}')
    return data


def borrar_checkpoint_pass1():
    if PASS1_CHECKPOINT.exists():
        PASS1_CHECKPOINT.unlink()
        print(f'  🗑️ Checkpoint eliminado: {PASS1_CHECKPOINT}')
    else:
        print('  ℹ️ No había checkpoint para borrar.')


# ─────────────────────────────────────────────────────────────────────────────
# 8.  EJECUCIÓN PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────

def ejecutar_pipeline_todos_los_ejes() -> None:
    print('=' * 70)
    print('  PIPELINE MECUN v4.6 — TODOS LOS EJES')
    print('=' * 70)

    checkpoint = cargar_checkpoint_pass1()

    if checkpoint is not None:
        df_largo                 = checkpoint['df_largo']
        originales_por_eje       = checkpoint['originales_por_eje']
        copias_entrantes_por_eje = checkpoint['copias_entrantes_por_eje']
    else:
        df_largo = preprocesar_wide_a_long(DATA_SOURCE_PATH)
        if df_largo.empty:
            raise SystemExit('❌ Preprocesamiento sin resultados.')

        # ─── PASS 1 — detectar cross-eje para todos los ejes originales ──────
        originales_por_eje: dict[str, pd.DataFrame] = {}
        copias_entrantes_por_eje: dict[str, list[pd.DataFrame]] = {
            eje: [] for eje in NOMBRES_EJES
        }

        print('\n' + '=' * 70)
        print('PASS 1 — DETECCIÓN CROSS-EJE EN LOS 5 EJES')
        print('=' * 70)

        for eje in NOMBRES_EJES:
            df_base = df_largo[df_largo['Eje_Tematico'] == eje].copy().reset_index(drop=True)
            print(f'\nDetectando cross-eje para: {eje} ({len(df_base)} propuestas originales)')

            if df_base.empty:
                originales_por_eje[eje] = df_base
                continue

            df_proc, dfs_copias     = detectar_cross_eje_lote(df_base, eje)
            originales_por_eje[eje] = df_proc

            for eje_dest, df_cop in dfs_copias.items():
                if not df_cop.empty:
                    copias_entrantes_por_eje[eje_dest].append(df_cop)

        guardar_checkpoint_pass1(df_largo, originales_por_eje, copias_entrantes_por_eje)

    # ─── PASS 2 — construir corpus final por eje y generar reportes ──────────
    print('\n' + '=' * 70)
    print('PASS 2 — REPORTES FINALES DE LOS 5 EJES')
    print('=' * 70)

    for eje in NOMBRES_EJES:
        frames: list[pd.DataFrame] = []

        df_orig = originales_por_eje.get(eje)
        if df_orig is not None and not df_orig.empty:
            frames.append(df_orig)

        if copias_entrantes_por_eje[eje]:
            frames.extend(copias_entrantes_por_eje[eje])

        if not frames:
            print(f'\nEje sin propuestas: {eje}. Se omite.')
            continue

        if 'Es_Copia' in pd.concat(frames, ignore_index=True).columns:
            pd.concat(frames, ignore_index=True)['Es_Copia'] = pd.concat(frames, ignore_index=True)['Es_Copia'].fillna(False).astype(bool)
        else:
            pd.concat(frames, ignore_index=True)['Es_Copia'] = False

        if 'Eje_Origen_Copia' in pd.concat(frames, ignore_index=True).columns:
            pd.concat(frames, ignore_index=True)['Eje_Origen_Copia'] = pd.concat(frames, ignore_index=True)['Eje_Origen_Copia'].fillna('').astype(str)
        else:
            pd.concat(frames, ignore_index=True)['Eje_Origen_Copia'] = ''

        if 'Ejes_Adicionales' in pd.concat(frames, ignore_index=True).columns:
            pd.concat(frames, ignore_index=True)['Ejes_Adicionales'] = pd.concat(frames, ignore_index=True)['Ejes_Adicionales'].apply(_asegurar_lista_segura)
        else:
            pd.concat(frames, ignore_index=True)['Ejes_Adicionales'] = [[] for _ in range(len(pd.concat(frames, ignore_index=True)))]

        try:
            procesar_eje_completo(pd.concat(frames, ignore_index=True), eje)
        except Exception as e:
            print(f'\n❌ Error en eje: {eje}')
            print(f'   Motivo: {e}')
            continue

    # ─── Crear carpetas de salida para todos los ejes ─────────────────────────
    for eje in NOMBRES_EJES:
        crear_carpeta_salida(eje)
        print(f'✅ Carpeta lista: {crear_carpeta_salida(eje)}')

    print('\n✅ Todas las carpetas listas')
