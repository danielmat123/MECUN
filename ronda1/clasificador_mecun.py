# --- INSTALACIÓN DE DEPENDENCIAS ---
#!pip install -q python-docx google-generativeai pandas

# --- 1. LIBRERÍAS Y CONEXIÓN A DATOS ---
import os
import re
import time
import random
import pandas as pd
import google.generativeai as genai
from google.colab import drive
from docx import Document
from docx.shared import Pt, RGBColor

print("Conectando con Google Drive...")
drive.mount('/content/drive')

# Tu ruta exacta de datos finales
DATA_SOURCE_PATH = 'ARCHIVO DE DATOS DINAL'

try:
    df = pd.read_csv(DATA_SOURCE_PATH, encoding='utf-8', on_bad_lines='skip')
    print(f"¡Éxito! Se cargaron {len(df)} filas del repositorio maestro.\n")
except Exception as e:
    print(f"Error crítico al cargar los datos: {e}")
    df = None


# --- 2. CONFIGURACIÓN DEL MODELO Y DICCIONARIO DE EJES ---

# Configuración de IA
API_KEY = os.environ.get("GEMINI_API_KEY", "API_KEY") # Reemplaza 
MODEL_NAME = 'gemini-2.5-flash'
MARGEN_ERROR_SLOVIN = 0.10

genai.configure(api_key=API_KEY)
modelo_ia = genai.GenerativeModel(MODEL_NAME)

# DICCIONARIO MAESTRO: Mapea cada eje con sus categorías a priori
DICCIONARIO_CATEGORIAS = {
    'Composición, funciones y características de cuerpos colegiados.': [
        "Reestructuración y Tamaño de los Consejos (CSU, CA, Sedes)",
        "Ampliación de Representación Estamentaria (Nuevos cupos, minorías, etc.)",
        "Modificación de Funciones y Toma de Decisiones (Poderes, vetos)",
        "Transparencia, Publicidad y Rendición de Cuentas del Cuerpo Colegiado",
        "Descentralización y Autonomía de Instancias Locales (Facultades/Sedes)",
        "Régimen de Inhabilidades, Incompatibilidades y Periodos",
        "Casos Aislados / Otros",
        "Órganos Consultivos Externos",
        "Formación para la Gobernanza y Cultura Democrática"
    ],
    'Designación y elección de autoridades académicas.': [
        "Elección Directa y Voto Universal",
        "Voto Ponderado por Estamentos",
        "Modelos Mixtos (Consulta + Decisión del CSU)",
        "Evaluación por Mérito Académico / Concurso",
        "Reglas Transversales y Garantías Electorales (Censo, foros, voto en blanco)",
        "Casos Aislados / Otros",
        "Inclusión y Representación de Grupos Específicos",
        "Rendición de Cuentas y Mecanismos de Revocatoria",
        "Representación y Ponderación Territorial (Sedes)"
    ],
    'Formas y mecanismos de participación democrática.': [
        "Plebiscitos, Referendos y Consultas Populares Universitarias",
        "Cabildos Abiertos, Consejos y Asambleas",
        "Iniciativas Normativas y Presupuestos Participativos",
        "Mecanismos Digitales y Plataformas de Voto Electrónico",
        "Veedurías Ciudadanas, Auditorías y Control Social",
        "Casos Aislados / Otros",
        "Despolitización de la Autonomía y Neutralidad Institucional",
        "Transparencia y Acceso a la Información Pública Institucional"
    ],
    'Fortalecimiento de la cultura política democrática de la comunidad universitaria.': [
        "Cátedras, Currículo y Formación en Participación Ciudadana",
        "Creación de Espacios de Debate y Deliberación Pública",
        "Estrategias Institucionales de Comunicación y Sensibilización",
        "Fomento del Liderazgo Estudiantil, Profesoral y Administrativo",
        "Ética, Transparencia y Acceso a la Información Pública",
        "Garantías para la Movilización y el Ejercicio Político",
        "Casos Aislados / Otros",
        "Bienestar, Convivencia y Prevención de Violencias"
    ],
    'Reconocimiento y fortalecimiento de las organizaciones de la comunidad universitaria.': [
        "Reconocimiento Institucional y Apoyo a Colectivos Estudiantiles",
        "Fortalecimiento de Sindicatos y Asociaciones de Profesores/Trabajadores",
        "Garantías de Financiación, Presupuesto y Espacios Físicos",
        "Participación e Incidencia de las Organizaciones en Decisiones",
        "Redes de Egresados y su Vinculación con la Universidad",
        "Grupos Académicos, Semilleros y Sociedades Científicas",
        "Casos Aislados / Otros",
        "Grupos Artísticos, Culturales y Deportivos",
        "Recuperación de la Memoria y Legado Histórico/Cultural",
        "Modelos de Economía Solidaria y Autogestión",
        "Inclusión y Reconocimiento de Pueblos Étnicos y Diversidades",
        "Seguridad, Convivencia y Garantías de Derechos Fundamentales",
        "Integración de Epistemologías Diversas y Sistemas de Conocimiento",
        "Ampliación del Concepto de Comunidad Universitaria",
        "Vinculación Universidad-Comunidad y Extensión Social",
        "Políticas Específicas de Inclusión y Accesibilidad para Personas con Discapacidad",
        "Reconocimiento Académico y de Bienestar del Trabajo Organizativo Estudiantil",
        "Adecuación de la Programación Académica para la Participación Estudiantil",
        "Gobernanza Democrática Interna y Transparencia de las Organizaciones Universitarias",
        "Acceso Socioeconómico y Asequibilidad de la Educación Superior"
    ]
}

# ⚠️ EJE PARA PROCESAR 
EJE_ACTUAL = 'Reconocimiento y fortalecimiento de las organizaciones de la comunidad universitaria.'

# Extraemos las categorías base correspondientes automáticamente
CATEGORIAS_BASE = DICCIONARIO_CATEGORIAS.get(EJE_ACTUAL, [])


# --- 3. FUNCIONES DE PREPARACIÓN ---

def preparar_dataset_eje(dataframe: pd.DataFrame, nombre_eje: str) -> pd.DataFrame:
    """Limpia, filtra y enriquece el dataframe para un eje temático específico."""
    df_temp = dataframe.dropna(subset=['Eje Temático MECUN', 'Propuesta'])
    df_eje = df_temp[df_temp['Eje Temático MECUN'] == nombre_eje].copy()

    cols_meta = ['Título de la propuesta', 'Nivel normativo de la reforma', 'Palabras Clave (3 a 5 palabras)']
    for col in cols_meta:
        df_eje[col] = df_eje.get(col, pd.Series(['No especificado'] * len(df_eje))).fillna('No especificado')

    df_eje['Texto_Enriquecido'] = (
        "Título: " + df_eje['Título de la propuesta'].astype(str) +
        " | Normativa: " + df_eje['Nivel normativo de la reforma'].astype(str) +
        " | Palabras Clave: " + df_eje['Palabras Clave (3 a 5 palabras)'].astype(str) +
        " | Propuesta: " + df_eje['Propuesta'].astype(str)
    )
    return df_eje

def calcular_muestra_estadistica(total: int, error: float) -> int:
    """Aplica la fórmula de Slovin con límites operativos."""
    if total <= 0: return 0
    n = int(total / (1 + total * (error**2)))
    return max(min(n, 50), 5)

if df is not None:
    df_eje = preparar_dataset_eje(df, EJE_ACTUAL)
    print(f"Dataset listo: {len(df_eje)} propuestas del eje seleccionado.\n")

    if df_eje.empty:
        print("❌ ERROR: El dataframe está vacío. Verifica que el nombre exacto de 'EJE_ACTUAL' coincida con la base de datos.")
        raise SystemExit("Ejecución detenida para evitar errores en la API.")
else:
    # Esto maneja el caso por si el archivo principal nunca cargó
    print("❌ ERROR CRÍTICO: No hay datos base (df es None). Revisa la carga de Drive.")
    raise SystemExit("Ejecución detenida.")

# --- 4. FASE DE CATEGORIZACIÓN DINÁMICA (Descubrimiento) ---

tamano_muestra = calcular_muestra_estadistica(len(df_eje), MARGEN_ERROR_SLOVIN)
print(f"Fase 1: Muestreo estadístico -> Población {len(df_eje)} | Muestra {tamano_muestra}")

textos_muestra = df_eje['Texto_Enriquecido'].sample(n=min(tamano_muestra, len(df_eje))).tolist()
bloque_textos_muestra = "\n---\n".join(textos_muestra)

prompt_descubrimiento = f"""
Eres un analista de datos evaluando propuestas universitarias.
Aquí tienes las categorías temáticas que YA existen (Categorías Base):
{chr(10).join(['- ' + c for c in CATEGORIAS_BASE])}

Lee la muestra de propuestas e identifica si surgen temas NUEVOS.

REGLAS ESTRICTAS DE CERO DUPLICIDAD Y SOBREAJUSTE:
1. Las nuevas categorías NO pueden ser sinónimos, parafraseos ni subtemas menores.
2. Antes de proponer una idea, pregúntate: "¿Esto podría encajar de forma lógica en alguna Categoría Base?". Si la respuesta es SÍ, NO la propongas.
3. Deben ser conceptos genuinamente distintos.
4. LA REGLA DE ORO: No crees una categoría para 1 sola propuesta aislada. Un tema nuevo debe ser un patrón claro.

Si TODAS las propuestas encajan en las Categorías Base, responde exactamente con la palabra "NINGUNA".
Si descubres PATRONES inéditos, responde ÚNICAMENTE con una lista enumerada.

Muestra de propuestas:
{bloque_textos_muestra}
"""

categorias_dinamicas = []
try:
    if tamano_muestra > 0:
        respuesta_descubrimiento = modelo_ia.generate_content(prompt_descubrimiento).text.strip()
        if respuesta_descubrimiento.upper() != "NINGUNA" and len(respuesta_descubrimiento) > 5:
            lineas = respuesta_descubrimiento.split('\n')
            for linea in lineas:
                texto_limpio = linea.lstrip('0123456789.-* ').strip()
                if texto_limpio:
                    categorias_dinamicas.append(texto_limpio)
            print(f"¡Genial! La IA descubrió {len(categorias_dinamicas)} categorías nuevas:")
            for c in categorias_dinamicas: print(f"  + {c}")
        else:
            print("La IA determinó que las categorías base son suficientes.")
except Exception as e:
    print(f"Hubo un error en el descubrimiento: {e}. Continuaremos con las bases.")

categorias_totales = CATEGORIAS_BASE + categorias_dinamicas + ["Casos Aislados / Otros"]
texto_categorias_clasificacion = "\n".join([f"{i+1}. {cat}" for i, cat in enumerate(categorias_totales)])


# --- 5. FASE DE CLASIFICACIÓN ESTRICTA ---
print("\nFase 2: Clasificando TODAS las propuestas y calculando nivel de confianza...")

prompt_clasificacion = f"""
Eres un jurado experto en normativa universitaria. Clasifica la propuesta ciudadana en UNA y SOLO UNA de las siguientes categorías exactas:

{texto_categorias_clasificacion}

Instrucciones estrictas:
- Responde con el formato exacto: "Nombre de la Categoría | Puntaje"
- El "Puntaje" debe ser un número del 1 al 100 indicando tu nivel de seguridad en esta clasificación (100 = certeza absoluta).
- No agregues explicaciones adicionales.

Contexto a clasificar:
{{texto}}
"""

etiquetas_asignadas = []
confianzas = []

for idx, texto in enumerate(df_eje['Texto_Enriquecido'].tolist()):
    instruccion = prompt_clasificacion.format(texto=texto)
    try:
        respuesta = modelo_ia.generate_content(instruccion).text.strip()

        if '|' in respuesta:
            partes = respuesta.split('|')
            etiqueta_cruda = partes[0].strip()
            puntaje = int(partes[1].strip())
        else:
            etiqueta_cruda = respuesta
            puntaje = 0

        etiqueta_limpia = etiqueta_cruda.lstrip('0123456789.- ').strip()

        etiquetas_asignadas.append(etiqueta_limpia)
        confianzas.append(puntaje)
        print(f"  [{idx+1}/{len(df_eje)}] Procesada. (Confianza: {puntaje}%)")

    except Exception as e:
        print(f"  [{idx+1}/{len(df_eje)}] Error de API, enviando a revisión manual.")
        etiquetas_asignadas.append("Casos Aislados / Otros")
        confianzas.append(0)

    #API (Rate Limit)
    time.sleep(1.5)

df_eje['Categoria_Final'] = etiquetas_asignadas
df_eje['Confianza'] = confianzas

# --- 6. ENSAMBLAJE DE DOCUMENTOS ENTREGABLES (Word y Excel) ---
print("\nFase 3: Sub-agrupando por ideas centrales y exportando Word/Excel...")

doc = Document()
doc.add_heading('Revisión de Propuestas MECUN', 0)
doc.add_heading(f'Eje Temático: {EJE_ACTUAL}', 1)

datos_excel = []
UMBRAL_CONFIANZA = 70
categorias_unicas = sorted(df_eje['Categoria_Final'].unique())

for idx_cat, categoria in enumerate(categorias_unicas):
    df_grupo = df_eje[df_eje['Categoria_Final'] == categoria]

    if idx_cat > 0:
        doc.add_page_break()

    doc.add_heading(f'Categoría Principal: {categoria}', level=1)
    doc.add_paragraph(f"Total de propuestas en esta categoría: {len(df_grupo)}").italic = True
    doc.add_paragraph("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")

    textos_numerados = []
    mapeo_propuestas = {}

    for idx_interno, (idx_real, fila) in enumerate(df_grupo.iterrows()):
        id_str = str(idx_interno + 1)
        texto_original = fila['Propuesta']
        textos_numerados.append(f"[{id_str}] {texto_original}")
        mapeo_propuestas[id_str] = {
            'texto': texto_original,
            'confianza': fila['Confianza'],
            'titulo': fila.get('Título de la propuesta', 'N/A')
        }

    bloque_textos = "\n".join(textos_numerados)

    prompt_agrupamiento = f"""
    Eres un analista normativo. Aquí tienes {len(df_grupo)} propuestas de la categoría '{categoria}'.
    Tu objetivo es REDUCIR EL TIEMPO DE LECTURA agrupando las propuestas que comparten la MISMA IDEA CENTRAL o MECANISMO.

    REGLAS:
    1. AGRUPA SIN MIEDO: Si comparten el mecanismo de fondo, van en el mismo grupo.
    2. SÍNTESIS: Para cada grupo, redacta 1 sola viñeta (máximo 15 palabras) que resuma la idea principal.
    3. EXCEPCIÓN: Solo deja una propuesta sola si su idea es completamente aislada o incompatible.

    Responde ESTRICTAMENTE con este formato repetitivo:
    GRUPO: [Frase corta de síntesis]
    IDs: [id1, id2]
    ---

    Propuestas a analizar:
    {bloque_textos}
    """

    try:
        respuesta_ia = modelo_ia.generate_content(prompt_agrupamiento).text.strip()
        bloques_agrupados = respuesta_ia.split('---')
    except Exception as e:
        print(f"  [!] Error de API en categoría '{categoria}'.")
        bloques_agrupados = []

    time.sleep(2) # Pausa de seguridad

    ids_procesados_por_ia = set()

    for bloque in bloques_agrupados:
        lineas = [l.strip() for l in bloque.strip().split('\n') if l.strip()]
        if not lineas: continue

        sintesis = ""
        ids_grupo = []

        for linea in lineas:
            linea_limpia = linea.replace('*', '').strip()
            if linea_limpia.upper().startswith("GRUPO:"):
                sintesis = linea_limpia.split(":", 1)[1].strip()
            elif "ID" in linea_limpia.upper():
                ids_grupo = re.findall(r'\d+', linea_limpia)

        ids_validos = [i for i in ids_grupo if i in mapeo_propuestas]

        if sintesis and ids_validos:
            p_grupo = doc.add_paragraph()
            p_grupo.add_run("💡 IDEA CENTRAL: ").bold = True
            p_grupo.add_run(sintesis).bold = True

            p_check = doc.add_paragraph("[   ] Viable      [   ] No Viable      [   ] Requiere Ajuste")
            p_check.italic = True

            p_respaldo = doc.add_paragraph("Textos originales que apoyan esta idea:")
            p_respaldo.runs[0].font.size = Pt(9)

            for id_p in ids_validos:
                datos = mapeo_propuestas[id_p]

                p_detalle = doc.add_paragraph(style='List Bullet')
                p_detalle.paragraph_format.left_indent = Pt(36)

                if datos['confianza'] < UMBRAL_CONFIANZA:
                    alerta = p_detalle.add_run("[⚠️ Dudosa] ")
                    alerta.font.color.rgb = RGBColor(255, 0, 0)

                p_detalle.add_run(f"[{id_p}] {datos['texto']}")

                datos_excel.append({
                    'Eje Temático': EJE_ACTUAL,
                    'Categoría Asignada': categoria,
                    'Síntesis del Grupo (Idea Central)': sintesis,
                    'Título Original': datos['titulo'],
                    'Texto Propuesta Original': datos['texto'],
                    'Confianza Clasificación (%)': datos['confianza']
                })

            doc.add_paragraph()
            ids_procesados_por_ia.update(ids_validos)

    ids_faltantes = [id_str for id_str in mapeo_propuestas.keys() if id_str not in ids_procesados_por_ia]

    if ids_faltantes:
        p_aisladas = doc.add_paragraph()
        p_aisladas.add_run("📌 Propuestas con ideas únicas (Sin agrupar):").bold = True

        for id_p in ids_faltantes:
            datos = mapeo_propuestas[id_p]

            p_detalle = doc.add_paragraph(style='Normal')
            p_detalle.paragraph_format.left_indent = Pt(18)

            if datos['confianza'] < UMBRAL_CONFIANZA:
                alerta = p_detalle.add_run("[⚠️ Dudosa] ")
                alerta.font.color.rgb = RGBColor(255, 0, 0)

            p_detalle.add_run(f"[{id_p}] {datos['texto']}")

            p_check = doc.add_paragraph("[   ] Viable      [   ] No Viable      [   ] Requiere Ajuste\n")
            p_check.italic = True
            p_check.paragraph_format.left_indent = Pt(18)

            datos_excel.append({
                'Eje Temático': EJE_ACTUAL,
                'Categoría Asignada': categoria,
                'Síntesis del Grupo (Idea Central)': '[Idea Única - Sin Agrupar]',
                'Título Original': datos['titulo'],
                'Texto Propuesta Original': datos['texto'],
                'Confianza Clasificación (%)': datos['confianza']
            })

# Guardado de Word y Excel Base
nombre_base = EJE_ACTUAL[:20].strip().replace(" ", "_").replace(",", "")
nombre_word = f'Reporte_Paginado_{nombre_base}.docx'
doc.save(nombre_word)
print(f"  -> Archivo Word guardado como '{nombre_word}'")

df_final = pd.DataFrame(datos_excel)
nombre_csv = f'Reporte_Tabular_{nombre_base}.csv'
df_final.to_csv(nombre_csv, index=False, sep=';', encoding='utf-8-sig')
print(f"  -> Archivo Tabular guardado como '{nombre_csv}'")


# --- 7. FASE 4: ANÁLISIS DE CONSENSOS Y DISENSOS ---
print(f"\nIniciando Fase 4: Análisis de Consensos y Disensos para el eje: {EJE_ACTUAL}")

# Instanciamos el modelo PRO para razonamiento complejo
modelo_juez = genai.GenerativeModel('gemini-2.5-pro')

# Preparamos el documento entregable
doc_semaforo = Document()
doc_semaforo.add_heading('Mapa de Consensos y Disensos - MECUN', 0)
doc_semaforo.add_heading(f'Eje Temático: {EJE_ACTUAL}', 1)
doc_semaforo.add_paragraph("Este documento identifica los puntos de acuerdo y tensión basándose estrictamente en la lectura íntegra de las propuestas ciudadanas originales.\n")

for idx_cat, categoria in enumerate(categorias_unicas):
    df_grupo = df_eje[df_eje['Categoria_Final'] == categoria]

    # 1. FILTRO LÓGICO: Exclusión de categorías inválidas o muy pequeñas
    if "Casos Aislados" in categoria or "Ninguna" in categoria:
        print(f"Omitiendo '{categoria}' (Categoría de descarte).")
        continue

    if len(df_grupo) < 2:
        print(f"Omitiendo '{categoria}' (Solo 1 propuesta, imposible buscar consensos).")
        
        if idx_cat > 0:
             doc_semaforo.add_page_break()
        doc_semaforo.add_heading(f'Categoría: {categoria}', level=2)
        doc_semaforo.add_paragraph("Categoría omitida del análisis de consensos y disensos por contener únicamente 1 propuesta.").italic = True
        continue

    print(f"Analizando tensiones en: {categoria} ({len(df_grupo)} propuestas)...")

    # 2. EXTRACCIÓN DE TEXTOS ÍNTEGROS
    textos_originales = []
    for idx, texto in enumerate(df_grupo['Propuesta'].tolist()):
        textos_originales.append(f"Propuesta [{idx+1}]: {texto}")

    bloque_textos_juez = "\n\n".join(textos_originales)

    # 3. PROMPT PARA BUSCAR CONSENSOS Y DISENSOS
    prompt_juez = f"""
    Eres un analista experto en políticas públicas y resolución de conflictos de la Universidad Nacional de Colombia.
    Estás analizando propuestas ciudadanas para el proceso de reforma normativa (MECUN).

    CONTEXTO:
    - Eje Temático: {EJE_ACTUAL}
    - Categoría de Análisis: {categoria}
    - Número de propuestas a leer: {len(df_grupo)}

    TU TAREA:
    Lee los textos íntegros y originales de las propuestas provistas abajo. Tu objetivo es identificar exclusivamente los consensos y disensos reales ENTRE ESTAS PROPUESTAS.
    No hagas un análisis político teórico general; cíñete estrictamente a lo que piden los autores en los textos.

    REGLAS DE FORMATO (SEMÁFORO):
    Debes ser extremadamente conciso, directo al grano, pero sin perder ningún detalle técnico o normativo clave. Explica el "por qué" de cada punto basado en los textos.

    Responde ÚNICAMENTE usando esta estructura exacta:

    🟢 CONSENSOS MAYORITARIOS:
    - [Idea central de acuerdo] - [Breve explicación de por qué es un consenso y qué detalles exigen en común].

    🟡 PUNTOS DE NEGOCIACIÓN (Tensiones resolubles / Zonas Grises):
    - [Punto de debate] - [Breve explicación de las diferentes posturas o vacíos sobre CÓMO implementar la idea].

    🔴 DISENSOS IRRECONCILIABLES:
    - [Punto de choque frontal] - [Breve explicación de las posturas opuestas y mutuamente excluyentes encontradas en los textos. Si no hay disensos graves, escribe "No se detectaron disensos irreconciliables en esta categoría"].

    TEXTOS ORIGINALES DE LAS PROPUESTAS:
    {bloque_textos_juez}
    """

    try:
        respuesta_juez = modelo_juez.generate_content(prompt_juez).text.strip()

        # 4. ESCRITURA EN EL DOCUMENTO
        if idx_cat > 0:
            doc_semaforo.add_page_break()

        doc_semaforo.add_heading(f'Categoría: {categoria}', level=2)
        doc_semaforo.add_paragraph(f"Basado en el análisis de {len(df_grupo)} propuestas íntegras.").italic = True

        for linea in respuesta_juez.split('\n'):
            linea_limpia = linea.strip()
            if not linea_limpia:
                continue

            p = doc_semaforo.add_paragraph()
            if linea_limpia.startswith('🟢') or linea_limpia.startswith('🟡') or linea_limpia.startswith('🔴'):
                p.add_run(linea_limpia).bold = True
                p.style = 'Heading 3'
            elif linea_limpia.startswith('-'):
                p.style = 'List Bullet'
                partes = linea_limpia.split('-', 2)
                if len(partes) > 2:
                    p.add_run("- " + partes[1].strip() + " -").bold = True
                    p.add_run(" " + partes[2].strip())
                else:
                    p.add_run(linea_limpia)
            else:
                p.add_run(linea_limpia)

    except Exception as e:
        print(f"  [!] Error al procesar la categoría '{categoria}': {e}")
        doc_semaforo.add_paragraph(f"Error al generar el mapa para esta categoría.")

    time.sleep(3)

# Guardar el documento final del Semáforo
nombre_semaforo = f'Mapa_Semaforo_{nombre_base}.docx'
doc_semaforo.save(nombre_semaforo)

print(f"\n¡Análisis completado! El mapa de consensos y disensos se guardó como '{nombre_semaforo}'")
